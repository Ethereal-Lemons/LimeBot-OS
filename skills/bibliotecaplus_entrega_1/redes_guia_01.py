from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = 'Guerrero_Velasquez_Leonel_Antonio_Guia_Practica_01.docx'

def shade(cell, fill='D9EAF7'):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def add_page_number(p):
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run()
    for kind, text in [('begin', None), ('instrText', ' PAGE '), ('end', None)]:
        el = OxmlElement('w:fldChar' if kind != 'instrText' else 'w:instrText')
        if kind == 'instrText':
            el.set(qn('xml:space'), 'preserve')
            el.text = text
        else:
            el.set(qn('w:fldCharType'), kind)
        r._r.append(el)

def para(doc, text='', bold=False, align=None, indent=0.5, style='Normal'):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.line_spacing = 2
    if indent is not None:
        p.paragraph_format.first_line_indent = Inches(indent)
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing = 2
    p.add_run(text)

def placeholder(doc, title, instructions):
    heading(doc, title, 2)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1
    r = p.add_run('[PEGAR CAPTURA AQUÍ]\n\n' + instructions)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = 'Times New Roman'

def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    for n in range(1, 4):
        st = doc.styles[f'Heading {n}']
        st.font.name = 'Times New Roman'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        st.font.size = Pt(12)
        st.font.bold = True
        st.paragraph_format.line_spacing = 2
    for n in ['List Bullet', 'List Number']:
        st = doc.styles[n]
        st.font.name = 'Times New Roman'
        st.font.size = Pt(12)
        st.paragraph_format.line_spacing = 2
    add_page_number(sec.header.paragraphs[0])

    # Portada
    cover = [
        'UNIVERSIDAD FRANCISCO GAVIDIA',
        'FACULTAD DE INGENIERÍA Y SISTEMAS',
        '',
        'GUÍA PRÁCTICA No. 01',
        'TÉRMINOS BÁSICOS DE REDES DE COMPUTADORAS',
        '',
        'Asignatura: Redes de Computadoras',
        'Grupo: N01',
        '',
        'Estudiante:',
        'Leonel Antonio Guerrero Velásquez',
        'Carrera: Ingeniería en Ciencias de la Computación',
        'Carné: GV100223',
        '',
        'Docente: Ing. Carlos Heriberto Henríquez Fermán',
        '',
        '16 de julio de 2026',
    ]
    for text in cover:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)
        if text and (text.startswith('UNIVERSIDAD') or text.startswith('FACULTAD') or text.startswith('GUÍA') or text.startswith('TÉRMINOS')):
            r.bold = True
    doc.add_page_break()

    heading(doc, 'Objetivo', 1)
    para(doc, 'Conocer y comprender diferentes términos básicos relacionados con las redes de computadoras, así como utilizar herramientas de diagnóstico para identificar información de conectividad, rutas de red, dirección IP pública y ancho de banda disponible.')

    heading(doc, 'Introducción', 1)
    para(doc, 'Las redes de computadoras permiten la comunicación entre dispositivos y el intercambio de información mediante diferentes protocolos y servicios. Para comprender su funcionamiento es necesario conocer conceptos como dirección IP, máscara de red, gateway, DNS, paquetes, routers y ancho de banda.')
    para(doc, 'En esta guía se presentan definiciones de conceptos fundamentales y se proponen actividades prácticas utilizando comandos y servicios web. Las capturas solicitadas deberán agregarse con los resultados obtenidos desde la conexión a Internet utilizada por el estudiante.')

    heading(doc, 'Conceptos básicos de redes', 1)
    concepts = [
        ('Ipconfig', 'ipconfig es una herramienta de línea de comandos incluida en Windows que permite consultar y administrar la configuración de red de los adaptadores del equipo. Al ejecutarlo sin parámetros muestra información como la dirección IPv4, la máscara de subred, el gateway predeterminado y, dependiendo de la configuración, otros datos relacionados con IPv6. También permite renovar o liberar configuraciones DHCP mediante parámetros específicos.'),
        ('Máscara de red', 'La máscara de red es un valor que determina qué parte de una dirección IP identifica la red y qué parte identifica al dispositivo dentro de esa red. Por ejemplo, en una red IPv4 con máscara 255.255.255.0, los primeros tres octetos representan la red y el último octeto identifica a los hosts. La máscara permite determinar si dos dispositivos pertenecen a la misma red local.'),
        ('Gateway o puerta de enlace', 'El gateway, también llamado puerta de enlace predeterminada, es el dispositivo que permite que un equipo se comunique con otras redes. En una red doméstica normalmente corresponde al router. Cuando el destino no pertenece a la red local, el equipo envía los paquetes al gateway para que este los encamine hacia Internet u otra red.'),
        ('DNS', 'DNS significa Domain Name System o Sistema de Nombres de Dominio. Su función principal es traducir nombres fáciles de recordar, como www.google.com, en direcciones IP que los dispositivos puedan utilizar para establecer una comunicación. Sin DNS, los usuarios tendrían que memorizar las direcciones IP de los servicios que desean visitar.'),
        ('Traceroute y tracert', 'Traceroute es una herramienta de diagnóstico que muestra los saltos o routers por los que pasa un paquete hasta llegar a un destino. En Windows se utiliza el comando tracert, mientras que en otros sistemas operativos suele utilizarse traceroute. La herramienta ayuda a analizar la ruta, los tiempos de respuesta y los puntos donde podría existir una interrupción o retraso.'),
        ('Dirección IP pública', 'La dirección IP pública es la dirección que identifica una conexión frente a Internet. Generalmente es asignada por el proveedor de servicios de Internet y puede ser estática o dinámica. Esta dirección no debe confundirse con la IP privada, que identifica los dispositivos dentro de la red local.'),
        ('Ancho de banda', 'El ancho de banda es la capacidad máxima de transmisión de datos de una conexión durante un período determinado. Normalmente se expresa en bits por segundo, por ejemplo Mbps o Gbps. Una conexión con mayor ancho de banda puede transferir una mayor cantidad de datos, aunque la velocidad real también depende de la latencia, la congestión, el servidor y la calidad de la conexión.'),
        ('Paquete de red', 'Un paquete es una unidad de datos que se envía a través de una red. La información grande se divide en paquetes para facilitar su transmisión. Cada paquete contiene datos y encabezados con información de control, como las direcciones de origen y destino, el protocolo y otros datos necesarios para su entrega y reconstrucción.'),
        ('Router', 'Un router o enrutador es un dispositivo que conecta diferentes redes y decide por cuál ruta deben enviarse los paquetes. Utiliza tablas de enrutamiento y direcciones IP para entregar la información a su destino. En una red doméstica también suele proporcionar funciones como Wi-Fi, DHCP, NAT y firewall básico.'),
    ]
    for title, text in concepts:
        heading(doc, title, 2)
        para(doc, text)

    heading(doc, 'Actividad práctica: comando tracert', 1)
    para(doc, 'En una computadora con Windows se debe abrir el símbolo del sistema y ejecutar el siguiente comando:')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2
    r = p.add_run('tracert www.google.com'); r.bold = True; r.font.name = 'Courier New'; r.font.size = Pt(12)
    placeholder(doc, 'Captura del comando tracert', 'Ejecute el comando y pegue aquí una captura donde se observe el resultado completo o la mayor parte de los saltos.')
    heading(doc, '¿Para qué se puede utilizar tracert?', 2)
    para(doc, 'tracert se utiliza para conocer la ruta que siguen los paquetes desde el equipo local hasta un destino específico. Muestra cada salto intermedio, normalmente correspondiente a un router, junto con los tiempos aproximados de respuesta. Esta información permite detectar dónde puede presentarse una falla, una demora elevada o una pérdida de comunicación. También ayuda a diferenciar si el problema se encuentra en la red local, en el proveedor de Internet o en una red intermedia.')

    heading(doc, 'Dirección IP pública y ubicación aproximada', 1)
    para(doc, 'La dirección IP pública debe consultarse desde la conexión a Internet utilizada para realizar la práctica. El resultado puede variar según la red, el proveedor y el momento de la consulta.')
    placeholder(doc, 'Consulta de dirección IP pública', 'Ingrese a cualesmiip.com, consulte la IP pública y pegue aquí la captura correspondiente.')
    placeholder(doc, 'Ubicación aproximada de la IP', 'Ingrese a myip.es o a un servicio equivalente, amplíe la información y pegue aquí la captura de la ubicación aproximada.')
    para(doc, 'La geolocalización de una dirección IP es aproximada. Generalmente identifica la ubicación del proveedor o del nodo de conexión y no necesariamente la ubicación exacta del usuario.')

    heading(doc, 'Prueba de ancho de banda', 1)
    placeholder(doc, 'Resultado de Speedtest', 'Ingrese a speedtest.net, ejecute la prueba y pegue aquí la captura donde se observen el ping, la velocidad de descarga y la velocidad de carga.')
    para(doc, 'El ping representa el tiempo aproximado que tarda un paquete en viajar hasta un servidor y regresar. La velocidad de descarga indica la cantidad de datos que pueden recibirse por segundo, mientras que la velocidad de carga indica la cantidad de datos que pueden enviarse por segundo. Los resultados pueden cambiar según el servidor elegido, la congestión de la red y los dispositivos conectados.')

    heading(doc, 'Consulta de información en LACNIC WHOIS', 1)
    para(doc, 'WHOIS es un servicio que permite consultar información de registro asociada con recursos de Internet, como bloques de direcciones IP y sistemas autónomos. En LACNIC, la consulta puede mostrar el organismo o proveedor al que está asignado un bloque de direcciones, el país, el estado del registro y datos técnicos disponibles públicamente.')
    placeholder(doc, 'Información de la IP pública en LACNIC', 'Consulte su IP pública en la sección WHOIS de LACNIC y pegue aquí la información o captura obtenida. Omita datos personales que no sean necesarios para la práctica.')

    heading(doc, 'Conclusión', 1)
    para(doc, 'Los conceptos estudiados son fundamentales para comprender cómo se comunican los dispositivos dentro de una red y cómo acceden a servicios externos. Herramientas como ipconfig y tracert permiten observar la configuración local y analizar la ruta hacia un destino, mientras que los servicios de consulta de IP y Speedtest proporcionan información sobre la conexión a Internet.')
    para(doc, 'La práctica también permite relacionar la teoría con una situación real, identificando la dirección IP pública, la ubicación aproximada asociada, el proveedor de Internet y el rendimiento disponible en el momento de la prueba.')

    heading(doc, 'Referencias', 1)
    refs = [
        'Microsoft. (s. f.). ipconfig. Microsoft Learn. https://learn.microsoft.com/windows-server/administration/windows-commands/ipconfig',
        'Microsoft. (s. f.). tracert. Microsoft Learn. https://learn.microsoft.com/windows-server/administration/windows-commands/tracert',
        'Internet Engineering Task Force. (1987). Domain names—Concepts and facilities (RFC 1034). https://www.rfc-editor.org/rfc/rfc1034',
        'Internet Engineering Task Force. (1981). Internet protocol (RFC 791). https://www.rfc-editor.org/rfc/rfc791',
        'LACNIC. (s. f.). Registro de Internet para América Latina y el Caribe. https://www.lacnic.net/',
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 2
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.add_run(ref)

    doc.core_properties.title = 'Guía Práctica 01 - Redes de Computadoras'
    doc.core_properties.author = 'Leonel Antonio Guerrero Velásquez'
    doc.core_properties.subject = 'Términos básicos de redes de computadoras'
    doc.save(OUT)
    print(OUT)

if __name__ == '__main__':
    main()
