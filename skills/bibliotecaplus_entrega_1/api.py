from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from datetime import date

OUTPUT = 'BibliotecaPlus_Primera_Entrega_APA.docx'


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f'Heading {level}')
    p.paragraph_format.keep_with_next = True
    p.add_run(text)
    return p


def add_body(doc, text, first_line=True):
    p = doc.add_paragraph(style='Normal')
    if first_line:
        p.paragraph_format.first_line_indent = Inches(0.5)
    p.add_run(text)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, True)
        set_cell_shading(hdr.cells[i], 'D9EAF7')
        hdr.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()
    return table


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = doc.styles
    normal = styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)

    for level in range(1, 4):
        style = styles[f'Heading {level}']
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        style.font.bold = True
        style.font.size = Pt(12)
        style.font.color.rgb = None
        style.paragraph_format.line_spacing = 2
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        if level == 1:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        else:
            style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    for style_name in ['List Bullet', 'List Number']:
        st = styles[style_name]
        st.font.name = 'Times New Roman'
        st._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        st.font.size = Pt(12)
        st.paragraph_format.line_spacing = 2

    # Header with page numbers
    header = sec.header.paragraphs[0]
    header.paragraph_format.space_after = Pt(0)
    add_page_number(header)

    # Cover page
    for text in [
        'UNIVERSIDAD FRANCISCO GAVIDIA',
        'FACULTAD DE INGENIERÍA Y SISTEMAS',
        'ADMINISTRACIÓN DE BASES DE DATOS',
        '',
        'PRIMERA ENTREGA',
        'PROPUESTA DE SISTEMA DE GESTIÓN DE BIBLIOTECA',
        '',
        'BibliotecaPlus',
        '',
        'Presentado por:',
        'Leonel Antonio Guerrero Velásquez',
        'Carné: GV100223',
        '',
        'Docente:',
        'Ing. Luis Enrique Reyes Valencia',
        '',
        'Ciclo II – 2026',
        '16 de julio de 2026',
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 2
        p.add_run(text).font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(12)
        if text in ['UNIVERSIDAD FRANCISCO GAVIDIA', 'FACULTAD DE INGENIERÍA Y SISTEMAS', 'ADMINISTRACIÓN DE BASES DE DATOS', 'PRIMERA ENTREGA', 'PROPUESTA DE SISTEMA DE GESTIÓN DE BIBLIOTECA', 'BibliotecaPlus']:
            p.runs[0].bold = True
    doc.add_page_break()

    # Contents placeholder using Word TOC field
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Contenido')
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    toc = doc.add_paragraph()
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'TOC \\o "1-3" \\h \\z \\u')
    toc._p.append(fld)
    doc.add_paragraph('Nota: en Microsoft Word, haga clic derecho sobre el índice y seleccione “Actualizar campo” para mostrar los números de página.')
    doc.add_page_break()

    add_heading(doc, 'Resumen', 1)
    add_body(doc, 'BibliotecaPlus es una propuesta de sistema informático para administrar los libros, usuarios, préstamos, devoluciones, reservas y multas de una biblioteca. El proyecto busca centralizar la información y automatizar las operaciones principales mediante una base de datos relacional. La solución permitirá consultar la disponibilidad de los libros, controlar las fechas de devolución, generar reportes y mantener un historial confiable de las operaciones. Esta primera entrega presenta el problema, la justificación, los objetivos, el alcance, los requisitos preliminares y la propuesta tecnológica del sistema.')
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.add_run('Palabras clave: ').bold = True
    p.add_run('biblioteca, base de datos, préstamos, devoluciones, reservas, multas.')

    add_heading(doc, 'Introducción', 1)
    add_body(doc, 'La administración de una biblioteca requiere controlar grandes cantidades de información relacionada con libros, usuarios y operaciones de préstamo. Cuando estos procesos se realizan de forma manual o mediante archivos independientes, aumentan las posibilidades de cometer errores, perder información o dificultar la consulta de los recursos disponibles.')
    add_body(doc, 'BibliotecaPlus propone una solución centralizada para organizar la información bibliográfica y automatizar los procesos de préstamo, devolución, reserva y control de multas. El sistema será diseñado utilizando una base de datos relacional y una aplicación web que permita a los usuarios autorizados realizar sus funciones de manera segura y ordenada.')

    add_heading(doc, 'Descripción del proyecto', 1)
    add_body(doc, 'BibliotecaPlus será un sistema de gestión bibliotecaria orientado a una biblioteca escolar, universitaria o comunitaria. Permitirá registrar los libros, autores, editoriales, categorías, ejemplares y usuarios. También permitirá controlar los préstamos y devoluciones, verificar la disponibilidad de los libros, registrar reservas y generar multas cuando existan retrasos.')
    add_body(doc, 'La aplicación contará con distintos niveles de acceso, de acuerdo con el rol de cada usuario. El administrador tendrá acceso a la configuración general y a los reportes; el bibliotecario gestionará las operaciones diarias; y el lector podrá consultar el catálogo, sus préstamos y sus reservas.')

    add_heading(doc, 'Planteamiento del problema', 1)
    add_body(doc, 'La biblioteca necesita un mecanismo eficiente para controlar sus recursos y las operaciones que se realizan diariamente. El manejo manual de los registros puede ocasionar inconsistencias entre la cantidad real de libros y la información registrada, además de dificultar el seguimiento de los préstamos y las devoluciones.')
    add_body(doc, 'Entre las principales dificultades identificadas se encuentran las siguientes:')
    for item in [
        'Dificultad para conocer qué libros están disponibles.',
        'Errores al registrar préstamos y devoluciones.',
        'Falta de control sobre las fechas límite de devolución.',
        'Cálculo manual e inconsistente de multas.',
        'Dificultad para consultar el historial de préstamos.',
        'Falta de seguimiento de las reservas pendientes.',
        'Demora en la atención de los usuarios.',
        'Ausencia de reportes para apoyar la toma de decisiones.'
    ]:
        add_bullet(doc, item)
    add_body(doc, 'Por esta razón, se propone desarrollar BibliotecaPlus como una herramienta que centralice la información, reduzca los errores y facilite la administración de la biblioteca.')

    add_heading(doc, 'Justificación', 1)
    add_body(doc, 'El desarrollo de BibliotecaPlus es importante porque permitirá mejorar la organización y el control de los recursos bibliográficos. Una base de datos centralizada facilitará la consulta de información y contribuirá a que los registros de libros, usuarios y préstamos sean consistentes.')
    add_body(doc, 'Desde el punto de vista académico, el proyecto permitirá aplicar conocimientos de administración de bases de datos mediante el diseño de tablas relacionadas, claves primarias y foráneas, restricciones de integridad, vistas, funciones, procedimientos almacenados, triggers y transacciones.')
    add_body(doc, 'Además, el sistema podrá servir como una solución inicial para instituciones que necesiten controlar sus operaciones bibliotecarias sin recurrir a procesos manuales. La propuesta mantiene un alcance realista y permite desarrollar una aplicación funcional dentro del período establecido.')

    add_heading(doc, 'Objetivos', 1)
    add_heading(doc, 'Objetivo general', 2)
    add_body(doc, 'Desarrollar un sistema informático para gestionar los libros, usuarios, préstamos, devoluciones, reservas y multas de una biblioteca, utilizando una base de datos relacional que garantice la integridad, disponibilidad y correcta administración de la información.')
    add_heading(doc, 'Objetivos específicos', 2)
    for item in [
        'Registrar y administrar la información de los libros disponibles en la biblioteca.',
        'Registrar los datos de los usuarios autorizados para solicitar préstamos.',
        'Controlar el proceso de préstamo y devolución de libros.',
        'Verificar automáticamente la disponibilidad de los libros.',
        'Registrar y administrar las reservas de libros.',
        'Calcular las multas generadas por devoluciones atrasadas.',
        'Mantener un historial de préstamos por usuario y por libro.',
        'Generar reportes sobre libros disponibles, préstamos vencidos y libros más solicitados.',
        'Implementar procedimientos, funciones, vistas y triggers para automatizar operaciones.',
        'Utilizar transacciones para garantizar que los préstamos y devoluciones se registren correctamente.'
    ]:
        add_bullet(doc, item)

    add_heading(doc, 'Usuarios del sistema', 1)
    add_table(doc, ['Rol', 'Descripción', 'Funciones principales'], [
        ['Administrador', 'Usuario con acceso general al sistema.', 'Gestionar usuarios, libros, parámetros, reportes y auditoría.'],
        ['Bibliotecario', 'Responsable de las operaciones diarias.', 'Registrar usuarios, préstamos, devoluciones, reservas y multas.'],
        ['Lector', 'Usuario registrado de la biblioteca.', 'Consultar catálogo, disponibilidad, préstamos, historial y reservas.'],
    ])

    add_heading(doc, 'Funcionalidades principales', 1)
    sections = [
        ('Gestión de usuarios', 'Permitirá registrar y administrar los datos de los usuarios de la biblioteca, incluyendo nombre completo, identificación, correo electrónico, teléfono, tipo de usuario, estado y fecha de registro.'),
        ('Gestión de libros', 'Permitirá registrar libros mediante datos como ISBN, título, autor, editorial, categoría, año de publicación, cantidad de ejemplares y estado.'),
        ('Gestión de autores, editoriales y categorías', 'Permitirá organizar la información bibliográfica y clasificar los libros por autor, editorial y temática.'),
        ('Registro de préstamos', 'Antes de registrar un préstamo, el sistema verificará que el usuario esté activo, que no tenga restricciones y que existan ejemplares disponibles.'),
        ('Registro de devoluciones', 'Al registrar una devolución, el sistema actualizará el préstamo, devolverá la disponibilidad del ejemplar y calculará una multa cuando corresponda.'),
        ('Gestión de reservas', 'Los usuarios podrán solicitar reservas para libros que no estén disponibles. Las reservas tendrán estados como pendiente, atendida, cancelada o vencida.'),
        ('Gestión de multas', 'El sistema calculará las multas con base en los días de atraso y permitirá controlar su estado de pago.'),
        ('Consultas y reportes', 'Se generarán reportes de libros disponibles, préstamos activos, préstamos vencidos, multas pendientes, reservas y libros más prestados.'),
    ]
    for title, text in sections:
        add_heading(doc, title, 2)
        add_body(doc, text)

    add_heading(doc, 'Requisitos funcionales preliminares', 1)
    add_table(doc, ['Código', 'Requisito'], [
        ['RF-01', 'El sistema deberá permitir iniciar sesión.'],
        ['RF-02', 'El sistema deberá manejar roles de usuario.'],
        ['RF-03', 'El administrador deberá gestionar usuarios del sistema.'],
        ['RF-04', 'El bibliotecario deberá registrar usuarios de la biblioteca.'],
        ['RF-05', 'El sistema deberá permitir registrar libros.'],
        ['RF-06', 'El sistema deberá permitir modificar y desactivar libros.'],
        ['RF-07', 'El sistema deberá permitir registrar autores, editoriales y categorías.'],
        ['RF-08', 'El sistema deberá permitir consultar la disponibilidad de los libros.'],
        ['RF-09', 'El sistema deberá registrar préstamos.'],
        ['RF-10', 'El sistema deberá registrar devoluciones.'],
        ['RF-11', 'El sistema deberá calcular multas por retraso.'],
        ['RF-12', 'El sistema deberá permitir registrar reservas.'],
        ['RF-13', 'El usuario deberá consultar sus préstamos activos.'],
        ['RF-14', 'El sistema deberá mostrar préstamos vencidos.'],
        ['RF-15', 'El sistema deberá generar reportes básicos.'],
        ['RF-16', 'El sistema deberá conservar el historial de préstamos.'],
    ])

    add_heading(doc, 'Requisitos no funcionales preliminares', 1)
    add_table(doc, ['Código', 'Requisito'], [
        ['RNF-01', 'La información deberá almacenarse en una base de datos relacional.'],
        ['RNF-02', 'El sistema deberá validar los datos ingresados.'],
        ['RNF-03', 'El sistema deberá impedir préstamos de libros no disponibles.'],
        ['RNF-04', 'El sistema deberá controlar el acceso según el rol.'],
        ['RNF-05', 'Las operaciones críticas deberán utilizar transacciones.'],
        ['RNF-06', 'La interfaz deberá ser clara y fácil de utilizar.'],
        ['RNF-07', 'El sistema deberá mantener la integridad referencial.'],
        ['RNF-08', 'Las contraseñas deberán almacenarse de forma segura.'],
        ['RNF-09', 'El sistema deberá responder adecuadamente ante errores.'],
        ['RNF-10', 'La base de datos deberá contar con un script de creación y respaldo.'],
    ])

    add_heading(doc, 'Propuesta preliminar de base de datos', 1)
    add_body(doc, 'Las principales entidades propuestas para la base de datos son UsuarioSistema, Rol, Lector, Libro, Autor, Editorial, Categoría, Ejemplar, Préstamo, Devolución, Reserva, Multa, Pago y Auditoría.')
    add_table(doc, ['Entidad', 'Descripción'], [
        ['UsuarioSistema', 'Almacena las credenciales y datos de acceso al sistema.'],
        ['Rol', 'Define los permisos de cada usuario del sistema.'],
        ['Lector', 'Almacena los datos de las personas que solicitan préstamos.'],
        ['Libro', 'Contiene los datos bibliográficos de cada título.'],
        ['Autor', 'Registra la información de los autores.'],
        ['Editorial', 'Registra las editoriales de los libros.'],
        ['Categoría', 'Clasifica los libros por temática.'],
        ['Ejemplar', 'Representa cada copia física de un libro.'],
        ['Préstamo', 'Registra la entrega temporal de un ejemplar a un lector.'],
        ['Devolución', 'Registra la devolución de un préstamo.'],
        ['Reserva', 'Controla las solicitudes de libros no disponibles.'],
        ['Multa', 'Registra los cargos por retrasos o pérdidas.'],
        ['Pago', 'Registra el pago de una multa.'],
        ['Auditoría', 'Conserva el historial de operaciones importantes.'],
    ])
    add_body(doc, 'Las relaciones principales serán las siguientes: un autor puede estar asociado con varios libros; una editorial puede publicar varios libros; una categoría puede contener varios libros; un libro puede tener varios ejemplares; un lector puede realizar muchos préstamos y reservas; y un préstamo puede generar una multa.')

    add_heading(doc, 'Objetos de base de datos propuestos', 1)
    add_heading(doc, 'Vistas', 2)
    for item in ['Vista de libros disponibles.', 'Vista de préstamos activos.', 'Vista de préstamos vencidos.', 'Vista de libros más prestados.', 'Vista de usuarios con multas pendientes.']:
        add_bullet(doc, item)
    add_heading(doc, 'Funciones almacenadas', 2)
    for item in ['Calcular días de atraso.', 'Calcular el monto de una multa.', 'Consultar la disponibilidad de un libro.', 'Determinar el límite de préstamos según el tipo de usuario.', 'Calcular el porcentaje de devoluciones realizadas a tiempo.']:
        add_bullet(doc, item)
    add_heading(doc, 'Procedimientos almacenados', 2)
    for item in ['Registrar préstamo.', 'Registrar devolución.', 'Registrar reserva.', 'Registrar pago de multa.', 'Registrar un nuevo libro.']:
        add_bullet(doc, item)
    add_heading(doc, 'Triggers', 2)
    for item in ['Actualizar la disponibilidad después de un préstamo.', 'Actualizar la disponibilidad después de una devolución.', 'Evitar que existan cantidades disponibles negativas.', 'Generar registros de auditoría.', 'Actualizar el estado de una reserva cuando un ejemplar vuelva a estar disponible.']:
        add_bullet(doc, item)

    add_heading(doc, 'Transacciones principales', 1)
    add_heading(doc, 'Transacción de préstamo', 2)
    add_body(doc, 'La transacción validará al usuario, verificará la disponibilidad del libro, registrará el préstamo y actualizará la cantidad de ejemplares disponibles. Si alguna operación falla, se ejecutará un rollback para evitar que la base de datos quede en un estado inconsistente.')
    add_heading(doc, 'Transacción de devolución', 2)
    add_body(doc, 'La transacción actualizará el préstamo, registrará la fecha de devolución, actualizará la disponibilidad, calculará la multa y registrará la deuda cuando corresponda. La operación será confirmada únicamente si todos los pasos se ejecutan correctamente.')

    add_heading(doc, 'Alcance del proyecto', 1)
    add_body(doc, 'El proyecto incluirá una aplicación web para gestionar usuarios, libros, autores, editoriales, categorías, préstamos, devoluciones, reservas y multas. También incluirá consultas, reportes básicos, control de acceso por roles, procedimientos, funciones, vistas, triggers, transacciones y documentación de instalación y uso.')

    add_heading(doc, 'Limitaciones del proyecto', 1)
    for item in ['Pagos en línea.', 'Aplicación móvil.', 'Integración con bibliotecas externas.', 'Lectura de códigos de barras o códigos QR.', 'Notificaciones por correo electrónico o SMS.', 'Digitalización de libros.', 'Préstamo de libros electrónicos.', 'Administración de múltiples sucursales.', 'Recomendaciones mediante inteligencia artificial.']:
        add_bullet(doc, item)
    add_body(doc, 'Estas funcionalidades podrán considerarse para una versión futura del sistema.')

    add_heading(doc, 'Tecnología propuesta', 1)
    add_table(doc, ['Componente', 'Tecnología propuesta'], [
        ['Lenguaje de programación', 'Java'],
        ['Framework', 'Spring Boot'],
        ['Base de datos', 'Microsoft SQL Server'],
        ['Frontend', 'HTML5, CSS3, JavaScript y Bootstrap'],
        ['Control de versiones', 'GitHub'],
        ['Modelado y diagramas', 'Draw.io'],
        ['Herramientas de desarrollo', 'Visual Studio Code o IntelliJ IDEA'],
        ['Administración de base de datos', 'SQL Server Management Studio'],
    ])
    add_body(doc, 'Java con Spring Boot ofrece una plataforma estable para desarrollar aplicaciones web y conectarlas con bases de datos relacionales. SQL Server permite implementar restricciones, consultas, vistas, funciones, procedimientos, triggers y transacciones. Bootstrap facilitará la elaboración de una interfaz clara y adaptable.')

    add_heading(doc, 'Plan de trabajo preliminar', 1)
    add_table(doc, ['Fase', 'Actividades principales'], [
        ['Análisis', 'Definición del problema, objetivos, alcance y requisitos.'],
        ['Diseño', 'Elaboración del modelo entidad-relación y diseño lógico.'],
        ['Base de datos', 'Creación de tablas, relaciones, restricciones y datos iniciales.'],
        ['Programación', 'Desarrollo de módulos, procedimientos, funciones y vistas.'],
        ['Integración', 'Conexión de la aplicación con la base de datos.'],
        ['Pruebas', 'Validación de préstamos, devoluciones, multas y reportes.'],
        ['Documentación', 'Elaboración del manual, informe y presentación final.'],
    ])

    add_heading(doc, 'Conclusión', 1)
    add_body(doc, 'BibliotecaPlus será una solución informática orientada a mejorar la administración de una biblioteca mediante el control de libros, usuarios, préstamos, devoluciones, reservas y multas. La propuesta es viable porque cuenta con un alcance controlado, utiliza tecnologías apropiadas y permite aplicar los principales conceptos de administración de bases de datos.')
    add_body(doc, 'El proyecto permitirá desarrollar una aplicación funcional para una situación real y demostrará el uso de tablas relacionadas, restricciones de integridad, vistas, funciones, procedimientos almacenados, triggers, transacciones y reportes. En las siguientes etapas se elaborará el modelo entidad-relación, el diseño lógico y la implementación de la base de datos.')

    add_heading(doc, 'Referencias', 1)
    refs = [
        'Microsoft. (s. f.). SQL Server documentation. Microsoft Learn. https://learn.microsoft.com/sql/sql-server/',
        'Oracle. (s. f.). Java documentation. https://docs.oracle.com/en/java/',
        'Spring. (s. f.). Spring Boot documentation. https://docs.spring.io/spring-boot/',
        'American Psychological Association. (2020). Publication manual of the American Psychological Association (7th ed.). American Psychological Association.'
    ]
    for ref in refs:
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.add_run(ref)

    # Set language and core properties
    doc.core_properties.title = 'BibliotecaPlus - Primera Entrega'
    doc.core_properties.author = 'Leonel Antonio Guerrero Velásquez'
    doc.core_properties.subject = 'Administración de Bases de Datos'
    doc.core_properties.comments = 'Documento elaborado con formato académico basado en APA 7.'
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == '__main__':
    main()
