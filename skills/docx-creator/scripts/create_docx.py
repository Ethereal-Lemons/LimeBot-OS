import sys
import argparse
from docx import Document


def create_document(title, filename, content_sections):
    document = Document()

    document.add_heading(title, 0)

    for section in content_sections:
        if section.get("heading"):
            document.add_heading(section["heading"], level=1)

        if section.get("text"):
            document.add_paragraph(section["text"])

    document.save(filename)
    print(f"Document saved as: {filename}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Create a Word document.")
    parser.add_argument("--title", required=True, help="Document title")
    parser.add_argument(
        "--filename", required=True, help="Output filename (e.g., homework.docx)"
    )
    parser.add_argument(
        "--content",
        nargs="+",
        required=True,
        help='Content in format "Heading:Text". Use multiple arguments for multiple sections.',
    )

    args = parser.parse_args()

    sections = []
    for item in args.content:
        if ":" in item:
            heading, text = item.split(":", 1)
            sections.append({"heading": heading, "text": text})
        else:
            sections.append({"heading": None, "text": item})

    try:
        create_document(args.title, args.filename, sections)
    except ImportError:
        print(
            "Error: 'python-docx' library is not installed. Run 'pip install python-docx' first."
        )
        sys.exit(1)
    except Exception as e:
        print(f"Error creating document: {e}")
        sys.exit(1)
